#!/usr/bin/env python3
"""
Debug específico para la función is_valid_docx
Identifica exactamente por qué falla la validación
"""
import sys
import tempfile
import zipfile
from pathlib import Path
from docx import Document

def debug_docx_validation():
    """Debug específico para validación DOCX"""
    print("🔍 DEBUG ESPECÍFICO PARA VALIDACIÓN DOCX")
    print("=" * 45)
    
    # Agregar backend al path
    sys.path.insert(0, 'backend/src')
    
    temp_dir = Path(tempfile.gettempdir()) / "docx_validation_debug"
    temp_dir.mkdir(exist_ok=True)
    
    # 1. Crear DOCX válido
    print("\n📄 1. CREANDO DOCX VÁLIDO...")
    try:
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('Simple test paragraph.')
        
        test_docx = temp_dir / "validation_test.docx"
        doc.save(str(test_docx))
        print(f"  ✅ DOCX creado: {test_docx}")
        
        # Verificar estructura manualmente
        with zipfile.ZipFile(test_docx, 'r') as zf:
            files = zf.namelist()
            print(f"  📋 Archivos en ZIP: {len(files)}")
            
            required = ['[Content_Types].xml', 'word/document.xml']
            print(f"  📋 Archivos requeridos:")
            for req in required:
                exists = req in files
                print(f"    {'✅' if exists else '❌'} {req}: {'EXISTE' if exists else 'FALTA'}")
            
            # Mostrar algunos archivos para debug
            print(f"  📋 Primeros 10 archivos:")
            for i, file in enumerate(files[:10]):
                print(f"    {i+1}. {file}")
                
    except Exception as e:
        print(f"  ❌ Error creando DOCX: {e}")
        return
    
    # 2. Probar función is_valid_docx directamente
    print("\n🔧 2. PROBANDO is_valid_docx DIRECTAMENTE...")
    try:
        from models.conversions.docx_to_html import is_valid_docx
        
        print(f"  • Probando is_valid_docx: ", end="")
        result = is_valid_docx(test_docx)
        print(f"{'✅ VÁLIDO' if result else '❌ INVÁLIDO'}")
        
        if not result:
            print("  🔍 Investigando por qué falló...")
            
            # Probar manualmente la lógica
            try:
                with zipfile.ZipFile(test_docx, 'r') as zip_file:
                    required_files = ['[Content_Types].xml', 'word/document.xml']
                    file_list = zip_file.namelist()
                    
                    print(f"    📋 Archivos en ZIP: {len(file_list)}")
                    print(f"    📋 Buscando archivos requeridos:")
                    
                    for required_file in required_files:
                        found = required_file in file_list
                        print(f"      {'✅' if found else '❌'} {required_file}: {'ENCONTRADO' if found else 'NO ENCONTRADO'}")
                        
                        if not found:
                            # Buscar archivos similares
                            similar = [f for f in file_list if required_file.lower() in f.lower()]
                            if similar:
                                print(f"        🔍 Archivos similares: {similar}")
                            
            except Exception as e:
                print(f"    ❌ Error en validación manual: {e}")
        
    except ImportError as e:
        print(f"  ❌ Error importando función: {e}")
    
    # 3. Probar función _read_docx directamente
    print("\n📖 3. PROBANDO _read_docx DIRECTAMENTE...")
    try:
        from models.conversions.docx_to_html import _read_docx
        
        print(f"  • Probando _read_docx: ", end="")
        text = _read_docx(test_docx)
        print(f"✅ ÉXITO - {len(text)} caracteres extraídos")
        print(f"  • Contenido: {text[:100]}...")
        
    except Exception as e:
        print(f"❌ ERROR: {e}")
        print(f"  🔍 Tipo de error: {type(e).__name__}")
    
    # 4. Probar Document directamente
    print("\n📚 4. PROBANDO Document DIRECTAMENTE...")
    try:
        doc = Document(test_docx)
        paragraphs = [p.text for p in doc.paragraphs]
        print(f"  ✅ Document cargado: {len(paragraphs)} párrafos")
        print(f"  📄 Contenido: {paragraphs}")
        
    except Exception as e:
        print(f"  ❌ Error con Document: {e}")
        print(f"  🔍 Tipo de error: {type(e).__name__}")
    
    # 5. Verificar si el problema es de encoding del archivo
    print("\n🔤 5. VERIFICANDO ENCODING DEL ARCHIVO...")
    try:
        from encoding_normalizer import detect_encoding, normalize_to_utf8
        
        # Leer bytes del archivo
        raw_bytes = test_docx.read_bytes()
        detected = detect_encoding(raw_bytes)
        print(f"  • Encoding detectado: {detected}")
        
        # Verificar si es un ZIP válido
        print(f"  • Es ZIP válido: ", end="")
        try:
            with zipfile.ZipFile(test_docx, 'r') as zf:
                zf.testzip()
            print("✅ SÍ")
        except Exception as e:
            print(f"❌ NO - {e}")
        
    except Exception as e:
        print(f"  ❌ Error verificando encoding: {e}")

def main():
    debug_docx_validation()

if __name__ == "__main__":
    main()
