#!/usr/bin/env python3
"""
Validación Final Completa del Sistema Anclora Nexus Mejorado
Prueba todas las mejoras implementadas: validación, encoding, mensajes, CSV, integración
"""
import os
import sys
import time
import json
import requests
import tempfile
from pathlib import Path
from docx import Document

# Configuración del servidor
BASE_URL = "http://localhost:8000"
API_BASE = f"{BASE_URL}/api"

class CompleteSystemValidator:
    """Validador completo del sistema mejorado"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "complete_validation"
        self.temp_dir.mkdir(exist_ok=True)
        self.results = {
            'encoding_system': {'passed': 0, 'failed': 0},
            'validation_system': {'passed': 0, 'failed': 0},
            'error_messages': {'passed': 0, 'failed': 0},
            'csv_optimization': {'passed': 0, 'failed': 0},
            'new_conversions': {'passed': 0, 'failed': 0},
            'overall': {'passed': 0, 'failed': 0}
        }

    def run_complete_validation(self):
        """Ejecuta validación completa del sistema"""
        print("🚀 VALIDACIÓN FINAL COMPLETA - SISTEMA ANCLORA NEXUS MEJORADO")
        print("=" * 70)
        
        # Verificar servidor
        if not self.verify_server():
            return
        
        # 1. Validar sistema de encoding (ya corregido)
        print("\n🔤 1. VALIDANDO SISTEMA DE ENCODING...")
        self.validate_encoding_system()
        
        # 2. Validar sistema de validación estricta
        print("\n🔒 2. VALIDANDO SISTEMA DE VALIDACIÓN ESTRICTA...")
        self.validate_validation_system()
        
        # 3. Validar mensajes de error mejorados
        print("\n💬 3. VALIDANDO MENSAJES DE ERROR MEJORADOS...")
        self.validate_error_messages()
        
        # 4. Validar optimización CSV
        print("\n📊 4. VALIDANDO OPTIMIZACIÓN CSV...")
        self.validate_csv_optimization()
        
        # 5. Validar nuevas conversiones
        print("\n🆕 5. VALIDANDO NUEVAS CONVERSIONES...")
        self.validate_new_conversions()
        
        # 6. Generar reporte final
        print("\n📋 6. GENERANDO REPORTE FINAL...")
        self.generate_final_report()

    def verify_server(self):
        """Verifica que el servidor esté funcionando"""
        try:
            response = requests.get(f"{BASE_URL}/api/health", timeout=5)
            if response.status_code == 200:
                print("✅ Servidor funcionando correctamente")
                return True
            else:
                print("❌ Servidor con problemas")
                return False
        except:
            print("❌ No se puede conectar al servidor")
            return False

    def validate_encoding_system(self):
        """Valida que el sistema de encoding funcione correctamente"""
        print("Validando sistema de encoding...")
        
        # Crear archivo con caracteres especiales
        test_file = self.temp_dir / "encoding_test.txt"
        test_content = "Texto con acentos: áéíóú ñÑ\nEmojis: 🚀 📄 ✅\nCaracteres especiales: © ® ™"
        test_file.write_text(test_content, encoding='utf-8')
        
        # Probar conversión TXT→HTML
        print("  • TXT con caracteres especiales → HTML: ", end="")
        success = self.test_conversion("encoding_test.txt", "txt", "html")
        
        if success:
            print("✅ ENCODING FUNCIONANDO")
            self.results['encoding_system']['passed'] += 1
        else:
            print("❌ ENCODING CON PROBLEMAS")
            self.results['encoding_system']['failed'] += 1

    def validate_validation_system(self):
        """Valida el sistema de validación estricta"""
        print("Validando sistema de validación estricta...")
        
        # Crear archivo corrupto
        corrupt_file = self.temp_dir / "corrupt.docx"
        corrupt_file.write_text("Este no es un DOCX real", encoding='utf-8')
        
        print("  • Archivo DOCX corrupto (debe ser rechazado): ", end="")
        success = self.test_conversion("corrupt.docx", "docx", "html")
        
        if not success:  # Debe fallar
            print("✅ RECHAZADO CORRECTAMENTE")
            self.results['validation_system']['passed'] += 1
        else:
            print("❌ ACEPTADO INCORRECTAMENTE")
            self.results['validation_system']['failed'] += 1
        
        # Crear DOCX válido
        print("  • Archivo DOCX válido (debe ser aceptado): ", end="")
        valid_docx = self.create_valid_docx()
        if valid_docx:
            success = self.test_conversion(valid_docx.name, "docx", "html")
            if success:
                print("✅ ACEPTADO CORRECTAMENTE")
                self.results['validation_system']['passed'] += 1
            else:
                print("❌ RECHAZADO INCORRECTAMENTE")
                self.results['validation_system']['failed'] += 1

    def validate_error_messages(self):
        """Valida que los mensajes de error sean amigables"""
        print("Validando mensajes de error mejorados...")
        
        # Probar con archivo corrupto para ver mensaje de error
        print("  • Mensaje de error para archivo corrupto: ", end="")
        
        corrupt_file = self.temp_dir / "error_test.docx"
        corrupt_file.write_bytes(b'archivo_corrupto_para_test')
        
        try:
            with open(corrupt_file, 'rb') as f:
                files = {'file': ('error_test.docx', f, 'application/octet-stream')}
                data = {'target_format': 'html'}
                
                response = requests.post(f"{API_BASE}/conversion/guest-convert", 
                                       files=files, data=data, timeout=15)
            
            if response.status_code == 400:  # Debe ser rechazado
                result = response.json()
                error_message = result.get('error', '')
                
                # Verificar que el mensaje sea amigable (no técnico)
                if 'ANX-' in error_message and 'Soluciones sugeridas' in error_message:
                    print("✅ MENSAJE AMIGABLE")
                    self.results['error_messages']['passed'] += 1
                else:
                    print("❌ MENSAJE TÉCNICO")
                    self.results['error_messages']['failed'] += 1
            else:
                print("❌ NO SE GENERÓ ERROR")
                self.results['error_messages']['failed'] += 1
                
        except Exception as e:
            print(f"❌ ERROR: {e}")
            self.results['error_messages']['failed'] += 1

    def validate_csv_optimization(self):
        """Valida las optimizaciones para CSV"""
        print("Validando optimización CSV...")
        
        # Crear CSV de prueba
        csv_file = self.temp_dir / "test_optimized.csv"
        csv_content = """Nombre,Edad,Ciudad,Salario
Juan Pérez,25,Madrid,45000.50
María García,30,Barcelona,52000.75
José López,28,Valencia,48000.00"""
        csv_file.write_text(csv_content, encoding='utf-8')
        
        # Probar preview CSV
        print("  • Preview CSV: ", end="")
        try:
            with open(csv_file, 'rb') as f:
                files = {'file': ('test_optimized.csv', f, 'text/csv')}
                
                response = requests.post(f"{API_BASE}/conversion/csv-preview", 
                                       files=files, timeout=15)
            
            if response.status_code == 200:
                result = response.json()
                if result.get('success') and 'analysis' in result:
                    analysis = result['analysis']
                    print(f"✅ PREVIEW GENERADO ({analysis['row_count']} filas)")
                    self.results['csv_optimization']['passed'] += 1
                else:
                    print("❌ PREVIEW INVÁLIDO")
                    self.results['csv_optimization']['failed'] += 1
            else:
                print(f"❌ ERROR HTTP {response.status_code}")
                self.results['csv_optimization']['failed'] += 1
                
        except Exception as e:
            print(f"❌ ERROR: {e}")
            self.results['csv_optimization']['failed'] += 1
        
        # Probar conversión CSV→HTML optimizada
        print("  • Conversión CSV→HTML optimizada: ", end="")
        success = self.test_conversion("test_optimized.csv", "csv", "html")
        
        if success:
            print("✅ CONVERSIÓN OPTIMIZADA")
            self.results['csv_optimization']['passed'] += 1
        else:
            print("❌ CONVERSIÓN FALLÓ")
            self.results['csv_optimization']['failed'] += 1

    def validate_new_conversions(self):
        """Valida las nuevas conversiones implementadas"""
        print("Validando nuevas conversiones...")
        
        new_conversions_to_test = [
            ('json', 'html', self.create_json_test),
            ('webp', 'jpg', self.create_webp_test),
        ]
        
        for source, target, creator_func in new_conversions_to_test:
            print(f"  • {source.upper()}→{target.upper()}: ", end="")
            
            try:
                test_file = creator_func()
                if test_file:
                    success = self.test_conversion(test_file.name, source, target)
                    if success:
                        print("✅ NUEVA CONVERSIÓN FUNCIONANDO")
                        self.results['new_conversions']['passed'] += 1
                    else:
                        print("❌ NUEVA CONVERSIÓN FALLÓ")
                        self.results['new_conversions']['failed'] += 1
                else:
                    print("⚠️ NO SE PUDO CREAR ARCHIVO TEST")
                    self.results['new_conversions']['failed'] += 1
            except Exception as e:
                print(f"❌ ERROR: {e}")
                self.results['new_conversions']['failed'] += 1

    def create_valid_docx(self):
        """Crea DOCX válido para pruebas"""
        try:
            doc = Document()
            doc.add_heading('Documento de Prueba', 0)
            doc.add_paragraph('Contenido de prueba simple.')
            
            docx_file = self.temp_dir / "valid_test.docx"
            doc.save(str(docx_file))
            return docx_file
        except:
            return None

    def create_json_test(self):
        """Crea archivo JSON para pruebas"""
        try:
            json_data = {
                "titulo": "Datos de Prueba",
                "items": [{"nombre": "Item 1", "valor": 100}]
            }
            json_file = self.temp_dir / "test.json"
            json_file.write_text(json.dumps(json_data, ensure_ascii=False), encoding='utf-8')
            return json_file
        except:
            return None

    def create_webp_test(self):
        """Crea archivo WEBP para pruebas (simulado)"""
        try:
            # Crear archivo que simula WEBP
            webp_file = self.temp_dir / "test.webp"
            webp_file.write_bytes(b'RIFF\x00\x00\x00\x00WEBP')  # Header WEBP básico
            return webp_file
        except:
            return None

    def test_conversion(self, filename, source_ext, target_ext):
        """Prueba una conversión específica"""
        try:
            file_path = self.temp_dir / filename
            if not file_path.exists():
                return False
            
            with open(file_path, 'rb') as f:
                files = {'file': (filename, f, 'application/octet-stream')}
                data = {'target_format': target_ext}
                
                response = requests.post(f"{API_BASE}/conversion/guest-convert", 
                                       files=files, data=data, timeout=30)
            
            return response.status_code == 200 and response.json().get('success', False)
            
        except Exception:
            return False

    def generate_final_report(self):
        """Genera reporte final completo"""
        print("\n" + "="*70)
        print("📊 REPORTE FINAL COMPLETO - SISTEMA ANCLORA NEXUS MEJORADO")
        print("="*70)
        
        # Calcular estadísticas totales
        total_passed = sum(cat['passed'] for cat in self.results.values())
        total_failed = sum(cat['failed'] for cat in self.results.values())
        total_tests = total_passed + total_failed
        overall_success = (total_passed / total_tests * 100) if total_tests > 0 else 0
        
        print(f"\n🎯 ESTADÍSTICAS GENERALES:")
        print(f"   Total de pruebas: {total_tests}")
        print(f"   Exitosas: {total_passed}")
        print(f"   Fallidas: {total_failed}")
        print(f"   Tasa de éxito general: {overall_success:.1f}%")
        
        # Detalles por categoría
        categories = {
            'encoding_system': '🔤 Sistema de Encoding',
            'validation_system': '🔒 Validación Estricta',
            'error_messages': '💬 Mensajes de Error',
            'csv_optimization': '📊 Optimización CSV',
            'new_conversions': '🆕 Nuevas Conversiones'
        }
        
        print(f"\n📋 RESULTADOS POR CATEGORÍA:")
        for key, title in categories.items():
            if key in self.results:
                cat_data = self.results[key]
                cat_total = cat_data['passed'] + cat_data['failed']
                cat_success = (cat_data['passed'] / cat_total * 100) if cat_total > 0 else 0
                status_icon = "✅" if cat_success >= 80 else "⚠️" if cat_success >= 60 else "❌"
                print(f"   {status_icon} {title}: {cat_success:.1f}% ({cat_data['passed']}/{cat_total})")
        
        # Evaluación del objetivo principal
        print(f"\n🎯 EVALUACIÓN DEL OBJETIVO:")
        if overall_success >= 90:
            print(f"   ✅ OBJETIVO ALCANZADO: Tasa de éxito ≥ 90%")
            print(f"   🚀 Sistema listo para producción")
        elif overall_success >= 80:
            print(f"   ⚠️ OBJETIVO PARCIAL: Tasa de éxito ≥ 80%")
            print(f"   🔧 Necesita ajustes menores")
        else:
            print(f"   ❌ OBJETIVO NO ALCANZADO: Tasa de éxito < 80%")
            print(f"   🛠️ Requiere correcciones adicionales")
        
        # Análisis específico de mejoras implementadas
        print(f"\n🔍 ANÁLISIS DE MEJORAS IMPLEMENTADAS:")
        
        encoding_success = self.results['encoding_system']['passed'] > 0
        validation_success = self.results['validation_system']['passed'] > 0
        error_msg_success = self.results['error_messages']['passed'] > 0
        csv_success = self.results['csv_optimization']['passed'] > 0
        new_conv_success = self.results['new_conversions']['passed'] > 0
        
        improvements = [
            ("Sistema de Encoding UTF-8", encoding_success, "Normalización automática y detección de archivos binarios"),
            ("Validación Estricta", validation_success, "Magic numbers, headers y verificación de integridad"),
            ("Mensajes de Error", error_msg_success, "Códigos únicos y soluciones específicas"),
            ("Optimización CSV", csv_success, "Detección automática y preview de datos"),
            ("Nuevas Conversiones", new_conv_success, "Formatos expandidos y motor Pandoc")
        ]
        
        for improvement, success, description in improvements:
            icon = "✅" if success else "❌"
            print(f"   {icon} {improvement}: {description}")
        
        # Recomendaciones finales
        print(f"\n💡 RECOMENDACIONES FINALES:")
        
        if overall_success >= 90:
            print("   • ✅ Sistema completamente validado y listo")
            print("   • 🚀 Proceder con despliegue en producción")
            print("   • 📊 Implementar monitoreo continuo")
            print("   • 🔄 Considerar expansión a más formatos")
        else:
            print("   • 🔧 Revisar categorías con fallos")
            print("   • 🛠️ Implementar correcciones específicas")
            print("   • 🧪 Ejecutar pruebas adicionales")
        
        # Guardar reporte completo
        report_data = {
            'timestamp': time.time(),
            'overall_success_rate': overall_success,
            'category_results': self.results,
            'improvements_status': {
                'encoding_system': encoding_success,
                'validation_system': validation_success,
                'error_messages': error_msg_success,
                'csv_optimization': csv_success,
                'new_conversions': new_conv_success
            }
        }
        
        with open('complete_system_validation_report.json', 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)
        
        print(f"\n📄 Reporte completo guardado en: complete_system_validation_report.json")

def main():
    """Función principal"""
    validator = CompleteSystemValidator()
    validator.run_complete_validation()

if __name__ == "__main__":
    main()
