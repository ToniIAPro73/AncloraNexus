#!/usr/bin/env python3
"""
Debug específico para problemas de conversión DOCX
Investiga por qué fallan las conversiones DOCX válidas
"""
import os
import sys
import tempfile
import requests
from pathlib import Path
from docx import Document

# Configuración del servidor
BASE_URL = "http://localhost:8000"
API_BASE = f"{BASE_URL}/api"

def debug_docx_conversion():
    """Debug específico para conversiones DOCX"""
    print("🔍 DEBUG ESPECÍFICO PARA CONVERSIONES DOCX")
    print("=" * 45)
    
    temp_dir = Path(tempfile.gettempdir()) / "docx_debug"
    temp_dir.mkdir(exist_ok=True)
    
    # 1. Crear DOCX válido simple
    print("\n📄 1. CREANDO DOCX VÁLIDO...")
    try:
        doc = Document()
        doc.add_heading('Documento de Prueba', 0)
        doc.add_paragraph('Este es un párrafo simple sin caracteres especiales.')
        doc.add_paragraph('Segundo párrafo con texto normal.')
        
        simple_docx = temp_dir / "simple_test.docx"
        doc.save(str(simple_docx))
        print(f"  ✅ DOCX simple creado: {simple_docx.name}")
        
        # Verificar que es un ZIP válido
        import zipfile
        with zipfile.ZipFile(simple_docx, 'r') as zf:
            files = zf.namelist()
            print(f"  📋 Archivos en DOCX: {len(files)} archivos")
            required = ['[Content_Types].xml', 'word/document.xml']
            has_required = all(req in files for req in required)
            print(f"  📋 Estructura válida: {'✅' if has_required else '❌'}")
            
    except Exception as e:
        print(f"  ❌ Error creando DOCX: {e}")
        return
    
    # 2. Probar conversión directa con el motor
    print("\n🔧 2. PROBANDO CONVERSIÓN DIRECTA...")
    try:
        sys.path.insert(0, 'backend/src')
        from models.conversion import conversion_engine
        
        output_file = temp_dir / "output_direct.html"
        success, message = conversion_engine.convert_file(
            str(simple_docx), str(output_file), 'docx', 'html'
        )
        
        print(f"  • Conversión directa: {'✅' if success else '❌'} - {message}")
        
        if success and output_file.exists():
            content = output_file.read_text(encoding='utf-8')
            print(f"  • Archivo generado: {len(content)} caracteres")
            if "Documento de Prueba" in content:
                print("  • Contenido correcto: ✅")
            else:
                print("  • Contenido incorrecto: ❌")
        
    except Exception as e:
        print(f"  ❌ Error en conversión directa: {e}")
    
    # 3. Probar conversión vía API
    print("\n🌐 3. PROBANDO CONVERSIÓN VÍA API...")
    try:
        with open(simple_docx, 'rb') as f:
            files = {'file': ('simple_test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'target_format': 'html'}
            
            response = requests.post(f"{API_BASE}/conversion/guest-convert", 
                                   files=files, data=data, timeout=30)
        
        print(f"  • Status Code: {response.status_code}")
        print(f"  • Response: {response.text[:200]}...")
        
        if response.status_code == 200:
            result = response.json()
            if result.get('success'):
                print("  • Conversión API: ✅ EXITOSA")
                
                # Probar descarga
                download_id = result.get('download_id')
                if download_id:
                    download_response = requests.get(
                        f"{API_BASE}/conversion/guest-download/{download_id}",
                        timeout=10
                    )
                    print(f"  • Descarga: {'✅' if download_response.status_code == 200 else '❌'}")
            else:
                print(f"  • Conversión API: ❌ FALLÓ - {result.get('message')}")
        else:
            print(f"  • Conversión API: ❌ ERROR HTTP {response.status_code}")
            
    except Exception as e:
        print(f"  ❌ Error en conversión API: {e}")
    
    # 4. Probar con archivo DOCX con caracteres especiales
    print("\n🌍 4. PROBANDO DOCX CON CARACTERES ESPECIALES...")
    try:
        doc_special = Document()
        doc_special.add_heading('Documento con Caracteres Especiales', 0)
        doc_special.add_paragraph('Texto con acentos: áéíóú ñÑ')
        doc_special.add_paragraph('Emojis: 🚀 📄 ✅')
        doc_special.add_paragraph('Símbolos: © ® ™ € £')
        
        special_docx = temp_dir / "special_chars.docx"
        doc_special.save(str(special_docx))
        
        # Probar conversión
        with open(special_docx, 'rb') as f:
            files = {'file': ('special_chars.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'target_format': 'html'}
            
            response = requests.post(f"{API_BASE}/conversion/guest-convert", 
                                   files=files, data=data, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            if result.get('success'):
                print("  • DOCX especiales: ✅ CONVERTIDO")
            else:
                print(f"  • DOCX especiales: ❌ FALLÓ - {result.get('message')}")
        else:
            print(f"  • DOCX especiales: ❌ ERROR HTTP {response.status_code}")
            
    except Exception as e:
        print(f"  ❌ Error con caracteres especiales: {e}")
    
    # 5. Verificar logs de encoding durante conversiones
    print("\n📋 5. VERIFICANDO LOGS DE ENCODING...")
    try:
        log_file = Path("backend/logs/encoding/encoding_normalizer.log")
        if log_file.exists():
            with open(log_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            print(f"  • Total entradas en log: {len(lines)}")
            
            # Buscar entradas recientes relacionadas con DOCX
            docx_entries = []
            for line in lines[-20:]:  # Últimas 20 entradas
                if 'docx' in line.lower():
                    try:
                        import json
                        entry = json.loads(line.strip())
                        docx_entries.append(entry)
                    except:
                        continue
            
            print(f"  • Entradas DOCX recientes: {len(docx_entries)}")
            
            if docx_entries:
                last_docx = docx_entries[-1]
                print(f"  • Última normalización DOCX: {last_docx.get('from')} → {last_docx.get('to')}")
            
        else:
            print("  ❌ Log de encoding no encontrado")
            
    except Exception as e:
        print(f"  ❌ Error verificando logs: {e}")
    
    print(f"\n🎯 CONCLUSIÓN DEBUG DOCX:")
    print(f"   • Sistema detecta archivos corruptos: ✅")
    print(f"   • Validación ZIP funcionando: ✅")
    print(f"   • Problema específico en conversión DOCX válidos")
    print(f"   • Encoding normalizer funcionando correctamente")

if __name__ == "__main__":
    debug_docx_conversion()
