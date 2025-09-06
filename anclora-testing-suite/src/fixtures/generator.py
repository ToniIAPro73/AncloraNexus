# src/fixtures/generator.py - Generador automático de fixtures para testing
import os
import random
import string
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Tuple
import logging
import json

# Imports para generación de archivos específicos
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
import pandas as pd
import io

logger = logging.getLogger(__name__)

class FixtureGenerator:
    """Generador master para todos los tipos de fixtures"""
    
    def __init__(self, fixtures_path: Path):
        self.fixtures_path = fixtures_path
        self.generated_files = {}
        self.generation_stats = {
            "total_generated": 0,
            "successful": 0,
            "failed": 0,
            "start_time": None,
            "end_time": None
        }
        
        # Contenido base para documentos
        self.lorem_ipsum = [
            "Lorem ipsum dolor sit amet, consectetur adipiscing elit.",
            "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.",
            "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris.",
            "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum.",
            "Excepteur sint occaecat cupidatat non proident, sunt in culpa qui.",
            "Officia deserunt mollit anim id est laborum consectetur adipiscing elit.",
            "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium.",
            "Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit."
        ]
        
        self.spanish_text = [
            "El sistema de testing de Anclora Nexus permite validar conversiones.",
            "Las secuencias automáticas son una ventaja competitiva única.",
            "La inteligencia artificial mejora la calidad de las conversiones.",
            "Los reportes en markdown son fáciles de leer y compartir.",
            "El análisis competitivo muestra nuestras fortalezas y oportunidades.",
            "La escalabilidad del sistema permite añadir nuevas conversiones.",
            "Los casos de prueba incluyen archivos corruptos y edge cases.",
            "La interfaz visual facilita el uso del sistema de testing."
        ]
        
        # Colores para imágenes
        self.colors = [
            (255, 0, 0),    # Rojo
            (0, 255, 0),    # Verde
            (0, 0, 255),    # Azul
            (255, 255, 0),  # Amarillo
            (255, 0, 255),  # Magenta
            (0, 255, 255),  # Cian
            (128, 128, 128), # Gris
            (255, 165, 0),  # Naranja
            (128, 0, 128),  # Púrpura
            (0, 128, 0),    # Verde oscuro
        ]
    
    def generate_all_fixtures(self) -> Dict[str, Any]:
        """Generar todos los fixtures necesarios para testing"""
        self.generation_stats["start_time"] = datetime.now()
        logger.info("🏭 Iniciando generación masiva de fixtures")
        
        try:
            # Crear directorios base
            self._create_directory_structure()
            
            # Generar fixtures por categoría
            self._generate_document_fixtures()    # ~290 archivos
            self._generate_image_fixtures()       # ~190 archivos
            self._generate_data_fixtures()        # ~45 archivos
            self._generate_media_fixtures()       # ~30 archivos
            self._generate_corrupted_fixtures()   # ~50 archivos
            self._generate_sequential_fixtures()  # ~115 archivos
            self._generate_edge_case_fixtures()   # ~40 archivos
            
            # Generar manifiesto
            self._generate_manifest()
            
            self.generation_stats["end_time"] = datetime.now()
            duration = (self.generation_stats["end_time"] - self.generation_stats["start_time"]).total_seconds()
            
            logger.info(f"✅ Generación completada: {self.generation_stats['successful']}/{self.generation_stats['total_generated']} en {duration:.1f}s")
            
            return {
                "stats": self.generation_stats,
                "generated_files": self.generated_files,
                "manifest_path": self.fixtures_path / "manifest.json"
            }
            
        except Exception as e:
            logger.error(f"❌ Error en generación de fixtures: {e}")
            raise
    
    def _create_directory_structure(self):
        """Crear estructura completa de directorios"""
        subdirs = [
            "documents", "images", "data", "media", 
            "corrupted", "sequential", "edge_cases"
        ]
        
        for subdir in subdirs:
            (self.fixtures_path / subdir).mkdir(parents=True, exist_ok=True)
            
        logger.info("📁 Estructura de directorios creada")
    
    def _generate_document_fixtures(self):
        """Generar fixtures de documentos (~290 archivos)"""
        logger.info("📄 Generando fixtures de documentos...")
        
        # TXT files (60 archivos)
        self._generate_txt_files()
        
        # DOCX files (50 archivos)
        self._generate_docx_files()
        
        # HTML files (40 archivos) 
        self._generate_html_files()
        
        # Markdown files (30 archivos)
        self._generate_md_files()
        
        # PDF files (20 archivos) - Simulados como placeholders
        self._generate_pdf_placeholders()
        
        # RTF files (15 archivos)
        self._generate_rtf_files()
        
        # Otros formatos (75 archivos)
        self._generate_other_document_formats()
    
    def _generate_txt_files(self):
        """Generar archivos TXT variados"""
        sizes = {"small": 1024, "medium": 51200, "large": 512000}  # 1KB, 50KB, 500KB
        encodings = ["utf-8", "latin-1", "cp1252"]
        content_types = ["simple", "structured", "multilingual", "special_chars"]
        
        for size_name, size_bytes in sizes.items():
            for encoding in encodings:
                for content_type in content_types:
                    for i in range(5):  # 5 variaciones cada uno
                        filename = f"txt_{size_name}_{encoding}_{content_type}_{i+1}.txt"
                        file_path = self.fixtures_path / "documents" / filename
                        
                        try:
                            content = self._generate_text_content(size_bytes, content_type)
                            with open(file_path, 'w', encoding=encoding, errors='replace') as f:
                                f.write(content)
                            
                            self._record_generated_file("documents", filename, file_path)
                            
                        except Exception as e:
                            logger.error(f"Error generando {filename}: {e}")
                            self._record_failed_file("documents", filename)
    
    def _generate_docx_files(self):
        """Generar archivos DOCX con diferentes características"""
        complexities = ["simple", "formatted", "complex"]
        features = ["text_only", "with_tables", "with_images_placeholder", "full_featured"]
        
        for complexity in complexities:
            for feature in features:
                for i in range(4):  # 4 variaciones
                    filename = f"docx_{complexity}_{feature}_{i+1}.docx"
                    file_path = self.fixtures_path / "documents" / filename
                    
                    try:
                        doc = Document()
                        self._populate_docx_document(doc, complexity, feature)
                        doc.save(file_path)
                        
                        self._record_generated_file("documents", filename, file_path)
                        
                    except Exception as e:
                        logger.error(f"Error generando {filename}: {e}")
                        self._record_failed_file("documents", filename)
        
        # Documentos adicionales con casos específicos (2 más)
        special_cases = ["empty_doc", "single_char"]
        for case in special_cases:
            filename = f"docx_special_{case}.docx"
            file_path = self.fixtures_path / "documents" / filename
            
            try:
                doc = Document()
                if case == "empty_doc":
                    pass  # Documento vacío
                elif case == "single_char":
                    doc.add_paragraph("A")
                    
                doc.save(file_path)
                self._record_generated_file("documents", filename, file_path)
                
            except Exception as e:
                logger.error(f"Error generando {filename}: {e}")
                self._record_failed_file("documents", filename)
    
    def _generate_html_files(self):
        """Generar archivos HTML variados"""
        html_types = ["simple", "styled", "complex", "broken"]
        
        for html_type in html_types:
            for i in range(10):  # 10 variaciones por tipo
                filename = f"html_{html_type}_{i+1}.html"
                file_path = self.fixtures_path / "documents" / filename
                
                try:
                    content = self._generate_html_content(html_type, i)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    self._record_generated_file("documents", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando {filename}: {e}")
                    self._record_failed_file("documents", filename)
    
    def _generate_md_files(self):
        """Generar archivos Markdown"""
        md_features = ["basic", "formatted", "advanced", "github_style"]
        
        for feature in md_features:
            for i in range(7):  # 7 archivos por feature (28 total)
                filename = f"md_{feature}_{i+1}.md"
                file_path = self.fixtures_path / "documents" / filename
                
                try:
                    content = self._generate_markdown_content(feature, i)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    self._record_generated_file("documents", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando {filename}: {e}")
                    self._record_failed_file("documents", filename)
        
        # 2 archivos Markdown adicionales especiales
        special_md = ["empty", "huge_table"]
        for special in special_md:
            filename = f"md_special_{special}.md"
            file_path = self.fixtures_path / "documents" / filename
            
            try:
                if special == "empty":
                    content = ""
                elif special == "huge_table":
                    content = self._generate_large_markdown_table()
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                self._record_generated_file("documents", filename, file_path)
                
            except Exception as e:
                logger.error(f"Error generando {filename}: {e}")
                self._record_failed_file("documents", filename)
    
    def _generate_image_fixtures(self):
        """Generar fixtures de imágenes (~190 archivos)"""
        logger.info("🖼️ Generando fixtures de imágenes...")
        
        # PNG files (50 archivos)
        self._generate_png_files()
        
        # JPEG files (40 archivos)
        self._generate_jpeg_files()
        
        # GIF files (30 archivos) - Básicos
        self._generate_gif_files()
        
        # SVG files (35 archivos)
        self._generate_svg_files()
        
        # WebP, TIFF y otros (35 archivos)
        self._generate_other_image_formats()
    
    def _generate_png_files(self):
        """Generar archivos PNG variados"""
        sizes = [(100, 100), (800, 600), (1920, 1080), (50, 50), (3000, 2000)]
        modes = ["RGB", "RGBA", "L", "P"]
        
        for i, size in enumerate(sizes):
            for j, mode in enumerate(modes):
                for variation in range(2):  # 2 variaciones por combinación
                    filename = f"png_{size[0]}x{size[1]}_{mode}_v{variation+1}.png"
                    file_path = self.fixtures_path / "images" / filename
                    
                    try:
                        img = self._create_test_image(size, mode, f"PNG Test {i}-{j}-{variation}")
                        img.save(file_path, "PNG")
                        
                        self._record_generated_file("images", filename, file_path)
                        
                    except Exception as e:
                        logger.error(f"Error generando {filename}: {e}")
                        self._record_failed_file("images", filename)
        
        # 10 PNG adicionales con características especiales
        special_pngs = [
            ("transparent", "RGBA", (400, 300)),
            ("grayscale", "L", (600, 400)),
            ("palette", "P", (300, 200)),
            ("tiny", "RGB", (10, 10)),
            ("square", "RGB", (500, 500))
        ]
        
        for special_type, mode, size in special_pngs:
            for i in range(2):
                filename = f"png_special_{special_type}_{i+1}.png"
                file_path = self.fixtures_path / "images" / filename
                
                try:
                    img = self._create_special_image(size, mode, special_type)
                    img.save(file_path, "PNG")
                    
                    self._record_generated_file("images", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando {filename}: {e}")
                    self._record_failed_file("images", filename)
    
    def _generate_svg_files(self):
        """Generar archivos SVG"""
        svg_types = ["simple", "complex", "with_text", "geometric", "gradients"]
        
        for svg_type in svg_types:
            for i in range(7):  # 7 archivos por tipo (35 total)
                filename = f"svg_{svg_type}_{i+1}.svg"
                file_path = self.fixtures_path / "images" / filename
                
                try:
                    content = self._generate_svg_content(svg_type, i)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    self._record_generated_file("images", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando {filename}: {e}")
                    self._record_failed_file("images", filename)
    
    def _generate_data_fixtures(self):
        """Generar fixtures de datos (~45 archivos)"""
        logger.info("📊 Generando fixtures de datos...")
        
        # CSV files (25 archivos)
        self._generate_csv_files()
        
        # JSON files (20 archivos)
        self._generate_json_files()
    
    def _generate_csv_files(self):
        """Generar archivos CSV variados"""
        csv_types = [
            ("simple", 10, 3),      # 10 filas, 3 columnas
            ("medium", 100, 5),     # 100 filas, 5 columnas  
            ("large", 1000, 10),    # 1000 filas, 10 columnas
            ("wide", 50, 20),       # 50 filas, 20 columnas
            ("mixed_types", 200, 8) # Tipos de datos mixtos
        ]
        
        for csv_type, rows, cols in csv_types:
            for i in range(5):  # 5 archivos por tipo
                filename = f"csv_{csv_type}_{rows}x{cols}_{i+1}.csv"
                file_path = self.fixtures_path / "data" / filename
                
                try:
                    df = self._generate_dataframe(csv_type, rows, cols)
                    df.to_csv(file_path, index=False, encoding='utf-8')
                    
                    self._record_generated_file("data", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando {filename}: {e}")
                    self._record_failed_file("data", filename)
    
    def _generate_json_files(self):
        """Generar archivos JSON variados"""
        json_types = ["simple", "nested", "array", "complex"]
        
        for json_type in json_types:
            for i in range(5):  # 5 archivos por tipo
                filename = f"json_{json_type}_{i+1}.json"
                file_path = self.fixtures_path / "data" / filename
                
                try:
                    data = self._generate_json_data(json_type, i)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(data, f, indent=2, ensure_ascii=False)
                    
                    self._record_generated_file("data", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando {filename}: {e}")
                    self._record_failed_file("data", filename)
    
    def _generate_corrupted_fixtures(self):
        """Generar archivos corruptos para testing de error handling (~50 archivos)"""
        logger.info("💥 Generando fixtures corruptos...")
        
        corruption_types = [
            "truncated",      # Archivo cortado
            "invalid_header", # Header inválido
            "random_bytes",   # Bytes aleatorios
            "encoding_error", # Error de encoding
            "mixed_content"   # Contenido mixto inválido
        ]
        
        formats = ["txt", "docx", "html", "csv", "json", "png", "svg"]
        
        for format_ext in formats:
            for corruption in corruption_types:
                filename = f"corrupted_{format_ext}_{corruption}.{format_ext}"
                file_path = self.fixtures_path / "corrupted" / filename
                
                try:
                    content = self._generate_corrupted_content(format_ext, corruption)
                    
                    # Escribir como binario para preservar corrupción
                    with open(file_path, 'wb') as f:
                        f.write(content)
                    
                    self._record_generated_file("corrupted", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando corrupted {filename}: {e}")
                    self._record_failed_file("corrupted", filename)
        
        # Casos especiales de corrupción (15 archivos adicionales)
        special_corruptions = [
            ("empty_file", "txt", b""),
            ("null_bytes", "txt", b"\x00" * 100),
            ("huge_line", "txt", b"A" * 1000000),
            ("binary_in_text", "html", b"<html>\x00\xFF\x01<body>test</body></html>"),
            ("invalid_utf8", "json", b'{"text": "\xFF\xFE invalid utf8"}')
        ]
        
        for special_name, ext, content in special_corruptions:
            for i in range(3):  # 3 variaciones de cada uno
                filename = f"corrupted_special_{special_name}_{i+1}.{ext}"
                file_path = self.fixtures_path / "corrupted" / filename
                
                try:
                    with open(file_path, 'wb') as f:
                        f.write(content)
                    
                    self._record_generated_file("corrupted", filename, file_path)
                    
                except Exception as e:
                    logger.error(f"Error generando special corrupted {filename}: {e}")
                    self._record_failed_file("corrupted", filename)
    
    def _generate_sequential_fixtures(self):
        """Generar fixtures para pruebas secuenciales (~115 archivos)"""
        logger.info("🔄 Generando fixtures para secuencias...")
        
        # Basado en el análisis competitivo de secuencias innovadoras
        sequence_types = [
            # Secuencias de 2 pasos
            ("docx_to_html_to_pdf", ["docx"], 20),
            ("svg_to_png_to_pdf", ["svg"], 15), 
            ("html_to_md_to_docx", ["html"], 15),
            
            # Secuencias de 3 pasos (críticas para competitividad)
            ("json_to_csv_to_xlsx_to_pdf", ["json"], 20),
            ("csv_to_json_to_html_to_pdf", ["csv"], 15),
            ("md_to_html_to_pdf_to_docx", ["md"], 10),
            
            # Secuencias bidireccionales
            ("docx_md_html_cycle", ["docx"], 10),
            ("csv_json_xlsx_cycle", ["csv"], 10)
        ]
        
        for seq_name, source_formats, count in sequence_types:
            for source_format in source_formats:
                for i in range(count):
                    filename = f"seq_{seq_name}_{source_format}_{i+1}.{source_format}"
                    file_path = self.fixtures_path / "sequential" / filename
                    
                    try:
                        # Generar archivo optimizado para secuencias
                        content = self._generate_sequential_content(source_format, seq_name)
                        
                        if source_format == "json":
                            with open(file_path, 'w', encoding='utf-8') as f:
                                json.dump(content, f, indent=2)
                        elif source_format == "csv":
                            content.to_csv(file_path, index=False)
                        elif source_format == "docx":
                            content.save(file_path)
                        elif source_format in ["html", "md", "svg"]:
                            with open(file_path, 'w', encoding='utf-8') as f:
                                f.write(content)
                        
                        self._record_generated_file("sequential", filename, file_path)
                        
                    except Exception as e:
                        logger.error(f"Error generando sequential {filename}: {e}")
                        self._record_failed_file("sequential", filename)
    
    # Métodos auxiliares para generación de contenido específico

    def _generate_special_chars_text(self, size_bytes: int) -> str:
        """Generar texto con caracteres especiales"""
        special_chars = "áéíóúñüç¿¡€£¥§©®™°±×÷≠≤≥∞∑∏∆√∫∂"
        base_text = "Texto con caracteres especiales: "

        while len(base_text.encode('utf-8')) < size_bytes:
            base_text += random.choice(special_chars) + " "
            if hasattr(self, 'lorem_ipsum'):
                base_text += random.choice(self.lorem_ipsum) + " "
            else:
                base_text += "lorem ipsum dolor sit amet "

        return base_text[:size_bytes]

    def _generate_markdown_content(self, feature: str, index: int) -> str:
        """Generar contenido markdown específico"""
        if feature == "github_style":
            return f"""# GitHub Style Markdown {index}

## Features
- [x] Task completed
- [ ] Task pending
- [x] Another completed task

### Code Block
```python
def hello_world():
    print("Hello from Anclora Nexus!")
```

### Table
| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Data 1   | Data 2   | Data 3   |
| Test A   | Test B   | Test C   |

> This is a blockquote for testing purposes.

**Bold text** and *italic text* for formatting tests.
"""
        else:
            return f"# Advanced Markdown {index}\n\nContent for advanced markdown testing."

    def _generate_large_markdown_table(self) -> str:
        """Generar tabla markdown grande"""
        headers = ["ID", "Name", "Email", "Department", "Salary", "Status"]
        header_line = "| " + " | ".join(headers) + " |"
        separator = "|" + "|".join(["---"] * len(headers)) + "|"

        rows = []
        for i in range(100):
            row_data = [
                str(i+1),
                f"Employee_{i+1}",
                f"emp{i+1}@company.com",
                random.choice(["IT", "HR", "Finance", "Marketing"]),
                f"${random.randint(30000, 100000)}",
                random.choice(["Active", "Inactive"])
            ]
            rows.append("| " + " | ".join(row_data) + " |")

        return "\n".join([header_line, separator] + rows)

    def _generate_pdf_placeholders(self):
        """Generar placeholders para archivos PDF"""
        logger.info("📄 Generando placeholders PDF...")
        # Por ahora solo crear archivos de texto que simulen PDFs
        for i in range(5):
            filename = f"pdf_placeholder_{i+1}.txt"
            content = f"PDF Placeholder {i+1} - This would be a PDF file in a real scenario"
            filepath = self.fixtures_path / "documents" / filename

            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
                self._record_generated_file("documents", filename, filepath)
            except Exception as e:
                logger.error(f"Error generando {filename}: {e}")
                self._record_failed_file("documents", filename)

    def _generate_text_content(self, size_bytes: int, content_type: str) -> str:
        """Generar contenido de texto según tipo y tamaño"""
        if content_type == "simple":
            base_text = " ".join(random.choices(self.lorem_ipsum, k=size_bytes // 50))
        elif content_type == "structured":
            base_text = self._generate_structured_text(size_bytes)
        elif content_type == "multilingual":
            base_text = self._generate_multilingual_text(size_bytes)
        elif content_type == "special_chars":
            base_text = self._generate_special_chars_text(size_bytes)
        else:
            base_text = " ".join(random.choices(self.lorem_ipsum, k=size_bytes // 50))
        
        # Ajustar al tamaño objetivo
        while len(base_text.encode('utf-8')) < size_bytes:
            base_text += " " + random.choice(self.lorem_ipsum)
        
        return base_text[:size_bytes]

    def _populate_docx_document(self, doc, complexity: str, feature: str):
        """Poblar documento DOCX según complejidad y características"""
        try:
            # Título principal
            title = doc.add_heading(f'Documento {complexity.title()} - {feature.replace("_", " ").title()}', 0)

            # Contenido según características
            if feature == "text_only":
                # Solo texto
                for i in range(3 if complexity == "simple" else 5 if complexity == "formatted" else 8):
                    p = doc.add_paragraph()
                    p.add_run(f"Párrafo {i+1}: ").bold = True
                    p.add_run(" ".join(random.choices(self.lorem_ipsum, k=2)))

            elif feature == "with_tables":
                # Texto + tablas
                doc.add_paragraph("Documento con tablas de ejemplo:")

                table = doc.add_table(rows=3, cols=3)
                table.style = 'Table Grid'

                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        cell.text = f"Celda {i+1},{j+1}"

            elif feature == "with_images_placeholder":
                # Texto + placeholders de imágenes
                doc.add_paragraph("Documento con placeholders de imágenes:")
                doc.add_paragraph("[IMAGEN PLACEHOLDER 1]")
                doc.add_paragraph("Texto entre imágenes.")
                doc.add_paragraph("[IMAGEN PLACEHOLDER 2]")

            elif feature == "full_featured":
                # Características completas
                doc.add_paragraph("Documento completo con múltiples características:")

                # Lista
                doc.add_paragraph("Lista de elementos:", style='List Bullet')
                for item in ["Elemento 1", "Elemento 2", "Elemento 3"]:
                    doc.add_paragraph(item, style='List Bullet')

                # Tabla
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'
                table.cell(0, 0).text = "Header 1"
                table.cell(0, 1).text = "Header 2"
                table.cell(1, 0).text = "Data 1"
                table.cell(1, 1).text = "Data 2"

        except Exception as e:
            logger.error(f"Error poblando documento DOCX: {e}")
            # Agregar contenido mínimo como fallback
            doc.add_paragraph(f"Documento de prueba {complexity} - {feature}")

    def _generate_html_content(self, html_type: str, index: int) -> str:
        """Generar contenido HTML según tipo"""
        try:
            if html_type == "simple":
                return f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>HTML Simple {index+1}</title>
</head>
<body>
    <h1>Documento HTML Simple {index+1}</h1>
    <p>Este es un documento HTML básico para testing.</p>
    <p>{random.choice(self.lorem_ipsum)}</p>
    <ul>
        <li>Elemento de lista 1</li>
        <li>Elemento de lista 2</li>
        <li>Elemento de lista 3</li>
    </ul>
</body>
</html>"""

            elif html_type == "formatted":
                return f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>HTML Formateado {index+1}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #333; }}
        .highlight {{ background-color: yellow; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; }}
    </style>
</head>
<body>
    <h1>Documento HTML Formateado {index+1}</h1>
    <p>Este documento incluye <span class="highlight">formato CSS</span>.</p>
    <table>
        <tr><th>Columna 1</th><th>Columna 2</th></tr>
        <tr><td>Dato 1</td><td>Dato 2</td></tr>
        <tr><td>Dato 3</td><td>Dato 4</td></tr>
    </table>
    <p><strong>Texto en negrita</strong> y <em>texto en cursiva</em>.</p>
</body>
</html>"""

            elif html_type == "complex":
                return f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>HTML Complejo {index+1}</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 0; padding: 20px; }}
        .container {{ max-width: 800px; margin: 0 auto; }}
        header {{ background: #f4f4f4; padding: 20px; margin-bottom: 20px; }}
        .grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }}
        .card {{ border: 1px solid #ddd; padding: 15px; border-radius: 5px; }}
        footer {{ margin-top: 40px; text-align: center; color: #666; }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>Documento HTML Complejo {index+1}</h1>
            <p>Ejemplo de documento con estructura avanzada</p>
        </header>

        <main>
            <div class="grid">
                <div class="card">
                    <h2>Sección 1</h2>
                    <p>{random.choice(self.lorem_ipsum)}</p>
                </div>
                <div class="card">
                    <h2>Sección 2</h2>
                    <p>{random.choice(self.lorem_ipsum)}</p>
                </div>
            </div>

            <section>
                <h2>Datos Tabulares</h2>
                <table style="width:100%; border-collapse: collapse;">
                    <thead>
                        <tr style="background: #f9f9f9;">
                            <th style="border: 1px solid #ddd; padding: 8px;">ID</th>
                            <th style="border: 1px solid #ddd; padding: 8px;">Nombre</th>
                            <th style="border: 1px solid #ddd; padding: 8px;">Valor</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr><td style="border: 1px solid #ddd; padding: 8px;">1</td><td style="border: 1px solid #ddd; padding: 8px;">Item A</td><td style="border: 1px solid #ddd; padding: 8px;">100</td></tr>
                        <tr><td style="border: 1px solid #ddd; padding: 8px;">2</td><td style="border: 1px solid #ddd; padding: 8px;">Item B</td><td style="border: 1px solid #ddd; padding: 8px;">200</td></tr>
                        <tr><td style="border: 1px solid #ddd; padding: 8px;">3</td><td style="border: 1px solid #ddd; padding: 8px;">Item C</td><td style="border: 1px solid #ddd; padding: 8px;">300</td></tr>
                    </tbody>
                </table>
            </section>
        </main>

        <footer>
            <p>Generado por Anclora Nexus Testing Suite</p>
        </footer>
    </div>
</body>
</html>"""

            else:
                # Fallback para tipos no reconocidos
                return f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>HTML Test {index+1}</title>
</head>
<body>
    <h1>Documento HTML de Prueba {index+1}</h1>
    <p>Tipo: {html_type}</p>
    <p>{random.choice(self.lorem_ipsum)}</p>
</body>
</html>"""

        except Exception as e:
            logger.error(f"Error generando contenido HTML: {e}")
            return f"""<!DOCTYPE html>
<html><head><title>Error</title></head>
<body><h1>Error generando contenido</h1><p>{str(e)}</p></body>
</html>"""
    
    def _generate_structured_text(self, size_bytes: int) -> str:
        """Generar texto estructurado con secciones"""
        sections = []
        current_size = 0
        section_num = 1
        
        while current_size < size_bytes:
            section_title = f"\n\nSECCIÓN {section_num}: TESTING ANCLORA NEXUS\n"
            section_content = "\n".join(random.choices(self.spanish_text, k=3))
            
            section = section_title + section_content
            sections.append(section)
            current_size += len(section.encode('utf-8'))
            section_num += 1
        
        return "\n".join(sections)
    
    def _generate_multilingual_text(self, size_bytes: int) -> str:
        """Generar texto multiidioma"""
        texts = {
            "es": self.spanish_text,
            "en": self.lorem_ipsum,
            "special": [
                "Testing with émojis: 🚀 💻 📊 ✅ ❌ ⚠️ 🎯",
                "Unicode characters: áéíóúñüç ¿¡ °±×÷ €£¥ ™®©",
                "Asian text: 你好世界 こんにちは Анклора Нексус العالم مرحبا",
                "Math symbols: ∑∫∏√∞≈≠≤≥±×÷∂∇∆Ω∴∵∈∉⊂⊃∪∩",
                "Arrows: ←↑→↓↔↕⟵⟶⟷⇒⇔⇄⇅↻↺⤴⤵⤶⤷",
            ]
        }
        
        combined_text = []
        current_size = 0
        
        while current_size < size_bytes:
            lang = random.choice(list(texts.keys()))
            text = random.choice(texts[lang])
            combined_text.append(text)
            current_size += len(text.encode('utf-8'))
        
        return " ".join(combined_text)
    
    def _create_test_image(self, size: Tuple[int, int], mode: str, text: str = "") -> Image.Image:
        """Crear imagen de prueba"""
        try:
            # Crear imagen base
            if mode == "RGBA":
                img = Image.new(mode, size, (255, 255, 255, 128))  # Semi-transparente
            elif mode == "L":
                img = Image.new(mode, size, 128)  # Gris medio
            elif mode == "P":
                img = Image.new("RGB", size, random.choice(self.colors))
                img = img.convert("P")
            else:
                img = Image.new(mode, size, random.choice(self.colors))
            
            draw = ImageDraw.Draw(img)
            width, height = size
            
            # Agregar elementos gráficos
            if width > 50 and height > 50:
                # Rectángulo
                rect_coords = [width//4, height//4, 3*width//4, 3*height//4]
                draw.rectangle(rect_coords, outline=random.choice(self.colors), width=2)
                
                # Círculo
                circle_coords = [width//3, height//3, 2*width//3, 2*height//3]
                if mode != "L":
                    draw.ellipse(circle_coords, fill=random.choice(self.colors))
                
                # Líneas
                draw.line([(0, 0), size], fill=random.choice(self.colors), width=2)
                draw.line([(0, height), (width, 0)], fill=random.choice(self.colors), width=2)
            
            # Agregar texto si es posible
            if text and width > 100 and height > 50:
                try:
                    draw.text((10, 10), text, fill=(0, 0, 0))
                except:
                    pass  # Ignorar si no hay fuente disponible
            
            return img
            
        except Exception as e:
            logger.error(f"Error creando imagen: {e}")
            # Fallback: imagen simple
            return Image.new("RGB", size, (255, 255, 255))
    
    def _record_generated_file(self, category: str, filename: str, file_path: Path):
        """Registrar archivo generado exitosamente"""
        if category not in self.generated_files:
            self.generated_files[category] = []
        
        file_info = {
            "filename": filename,
            "path": str(file_path),
            "size_bytes": file_path.stat().st_size if file_path.exists() else 0,
            "checksum": self._calculate_checksum(file_path),
            "generated_at": datetime.now().isoformat()
        }
        
        self.generated_files[category].append(file_info)
        self.generation_stats["successful"] += 1
        self.generation_stats["total_generated"] += 1
    
    def _record_failed_file(self, category: str, filename: str):
        """Registrar archivo que falló al generar"""
        self.generation_stats["failed"] += 1
        self.generation_stats["total_generated"] += 1
        logger.warning(f"❌ Falló generación: {category}/{filename}")
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calcular checksum SHA256 de archivo"""
        try:
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except:
            return "error_calculating_checksum"
    
    def _generate_manifest(self):
        """Generar manifiesto con información de todos los fixtures"""
        manifest = {
            "generator": "FixtureGenerator",
            "version": "1.0.0",
            "generated_at": datetime.now().isoformat(),
            "stats": self.generation_stats,
            "categories": {}
        }
        
        total_files = 0
        total_size = 0
        
        for category, files in self.generated_files.items():
            category_size = sum(f["size_bytes"] for f in files)
            manifest["categories"][category] = {
                "file_count": len(files),
                "total_size_bytes": category_size,
                "files": files
            }
            total_files += len(files)
            total_size += category_size
        
        manifest["summary"] = {
            "total_files": total_files,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2)
        }
        
        # Guardar manifiesto
        manifest_path = self.fixtures_path / "manifest.json"
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, indent=2, ensure_ascii=False)
        
        logger.info(f"📋 Manifiesto generado: {total_files} archivos, {manifest['summary']['total_size_mb']} MB")

# Funciones auxiliares para simplificar uso

def generate_all_fixtures(fixtures_path: Path = None) -> Dict[str, Any]:
    """Función de conveniencia para generar todos los fixtures"""
    if fixtures_path is None:
        fixtures_path = Path("./fixtures")
    
    generator = FixtureGenerator(fixtures_path)
    return generator.generate_all_fixtures()

if __name__ == "__main__":
    # Permitir ejecución directa
    import sys
    from pathlib import Path
    
    fixtures_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("./fixtures")
    
    logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
    
    result = generate_all_fixtures(fixtures_path)
    
    print("\n🎉 Generación de fixtures completada:")
    print(f"✅ Exitosos: {result['stats']['successful']}")
    print(f"❌ Fallidos: {result['stats']['failed']}")
    print(f"📁 Ubicación: {fixtures_path}")
    print(f"📋 Manifiesto: {result['manifest_path']}")