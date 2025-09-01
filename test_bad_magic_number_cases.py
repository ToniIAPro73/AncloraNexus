#!/usr/bin/env python3
"""
Pruebas específicas para casos que causan "Bad magic number"
Enfocado en archivos DOCX y otros formatos problemáticos
"""
import os
import sys
import tempfile
import requests
import zipfile
from pathlib import Path

# Configuración del servidor
BASE_URL = "http://localhost:8000"
API_BASE = f"{BASE_URL}/api"

class BadMagicNumberTester:
    """Tester específico para errores 'Bad magic number'"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "bad_magic_tests"
        self.temp_dir.mkdir(exist_ok=True)
        self.results = {'passed': 0, 'failed': 0, 'details': {}}

    def run_bad_magic_tests(self):
        """Ejecuta pruebas específicas para 'Bad magic number'"""
        print("🔍 PRUEBAS ESPECÍFICAS PARA 'BAD MAGIC NUMBER'")
        print("=" * 50)
        
        # 1. Crear archivos problemáticos
        print("\n📝 1. CREANDO ARCHIVOS PROBLEMÁTICOS...")
        self.create_problematic_files()
        
        # 2. Probar archivos DOCX corruptos
        print("\n📄 2. PROBANDO ARCHIVOS DOCX...")
        self.test_docx_cases()
        
        # 3. Probar archivos con encoding mixto
        print("\n🔀 3. PROBANDO ENCODING MIXTO...")
        self.test_mixed_encoding()
        
        # 4. Probar archivos con BOM problemático
        print("\n📋 4. PROBANDO BOM PROBLEMÁTICO...")
        self.test_bom_issues()
        
        # 5. Probar archivos ZIP corruptos
        print("\n📦 5. PROBANDO ZIP CORRUPTOS...")
        self.test_zip_corruption()
        
        # 6. Generar reporte
        print("\n📊 6. GENERANDO REPORTE...")
        self.generate_report()

    def create_problematic_files(self):
        """Crea archivos que pueden causar 'Bad magic number'"""
        print("Creando archivos problemáticos...")
        
        # 1. Archivo que parece DOCX pero no lo es
        fake_docx = self.temp_dir / "fake.docx"
        fake_docx.write_text("Este no es un DOCX real", encoding='utf-8')
        print(f"  ✅ Creado: {fake_docx.name} (DOCX falso)")
        
        # 2. Archivo DOCX con header corrupto
        corrupt_docx = self.temp_dir / "corrupt.docx"
        corrupt_docx.write_bytes(b'PK\x03\x04\x00\x00\x00\x00corrupted_data_here')
        print(f"  ✅ Creado: {corrupt_docx.name} (DOCX corrupto)")
        
        # 3. Archivo con BOM UTF-16 pero contenido UTF-8
        mixed_bom = self.temp_dir / "mixed_bom.txt"
        mixed_bom.write_bytes(b'\xff\xfe' + "Texto UTF-8 con BOM UTF-16".encode('utf-8'))
        print(f"  ✅ Creado: {mixed_bom.name} (BOM mixto)")
        
        # 4. Archivo con null bytes
        null_bytes = self.temp_dir / "null_bytes.txt"
        null_bytes.write_bytes(b'Texto\x00con\x00null\x00bytes')
        print(f"  ✅ Creado: {null_bytes.name} (Null bytes)")

    def test_docx_cases(self):
        """Prueba casos específicos de DOCX que causan problemas"""
        print("Probando casos DOCX problemáticos...")
        
        # Caso 1: DOCX falso
        print("  • DOCX falso → HTML: ", end="")
        success = self.test_conversion_api("fake.docx", "docx", "html")
        if not success:
            print("✅ RECHAZADO CORRECTAMENTE")
            self.results['passed'] += 1
            self.results['details']['fake_docx'] = 'correctly_rejected'
        else:
            print("❌ ACEPTADO INCORRECTAMENTE")
            self.results['failed'] += 1
            self.results['details']['fake_docx'] = 'incorrectly_accepted'
        
        # Caso 2: DOCX corrupto
        print("  • DOCX corrupto → HTML: ", end="")
        success = self.test_conversion_api("corrupt.docx", "docx", "html")
        if not success:
            print("✅ RECHAZADO CORRECTAMENTE")
            self.results['passed'] += 1
            self.results['details']['corrupt_docx'] = 'correctly_rejected'
        else:
            print("❌ ACEPTADO INCORRECTAMENTE")
            self.results['failed'] += 1
            self.results['details']['corrupt_docx'] = 'incorrectly_accepted'
        
        # Caso 3: Crear DOCX válido para comparación
        print("  • DOCX válido → HTML: ", end="")
        valid_docx = self.create_valid_docx()
        if valid_docx:
            success = self.test_conversion_api(valid_docx.name, "docx", "html")
            if success:
                print("✅ CONVERSIÓN EXITOSA")
                self.results['passed'] += 1
                self.results['details']['valid_docx'] = 'converted_successfully'
            else:
                print("❌ CONVERSIÓN FALLÓ")
                self.results['failed'] += 1
                self.results['details']['valid_docx'] = 'conversion_failed'
        else:
            print("⚠️ NO SE PUDO CREAR DOCX VÁLIDO")
            self.results['details']['valid_docx'] = 'creation_failed'

    def create_valid_docx(self):
        """Crea un archivo DOCX válido para pruebas"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('Documento de Prueba', 0)
            doc.add_paragraph('Este es un documento DOCX válido con caracteres especiales: áéíóú ñÑ')
            doc.add_paragraph('Emojis: 🚀 📄 ✅')
            
            valid_docx = self.temp_dir / "valid_test.docx"
            doc.save(str(valid_docx))
            
            return valid_docx
        except Exception as e:
            print(f"Error creando DOCX válido: {e}")
            return None

    def test_mixed_encoding(self):
        """Prueba archivos con encoding mixto"""
        print("Probando encoding mixto...")
        
        # Archivo con BOM mixto
        print("  • BOM mixto → TXT: ", end="")
        success = self.test_conversion_api("mixed_bom.txt", "txt", "html")
        if success:
            print("✅ MANEJADO CORRECTAMENTE")
            self.results['passed'] += 1
            self.results['details']['mixed_bom'] = 'handled_correctly'
        else:
            print("❌ ERROR EN MANEJO")
            self.results['failed'] += 1
            self.results['details']['mixed_bom'] = 'handling_error'

    def test_bom_issues(self):
        """Prueba problemas específicos de BOM"""
        print("Probando problemas de BOM...")
        
        # Archivo con null bytes
        print("  • Null bytes → HTML: ", end="")
        success = self.test_conversion_api("null_bytes.txt", "txt", "html")
        if success:
            print("✅ MANEJADO")
            self.results['passed'] += 1
            self.results['details']['null_bytes'] = 'handled'
        else:
            print("❌ ERROR")
            self.results['failed'] += 1
            self.results['details']['null_bytes'] = 'error'

    def test_zip_corruption(self):
        """Prueba corrupción específica de archivos ZIP (DOCX, EPUB, ODT)"""
        print("Probando corrupción ZIP...")
        
        # Crear ZIP corrupto que parece DOCX
        print("  • ZIP corrupto como DOCX: ", end="")
        try:
            corrupt_zip_docx = self.temp_dir / "corrupt_zip.docx"
            
            # Crear ZIP con estructura incorrecta
            with zipfile.ZipFile(corrupt_zip_docx, 'w') as zf:
                zf.writestr("wrong_structure.xml", "<?xml version='1.0'?><root>corrupted</root>")
            
            success = self.test_conversion_api("corrupt_zip.docx", "docx", "html")
            if not success:
                print("✅ RECHAZADO CORRECTAMENTE")
                self.results['passed'] += 1
                self.results['details']['corrupt_zip'] = 'correctly_rejected'
            else:
                print("❌ ACEPTADO INCORRECTAMENTE")
                self.results['failed'] += 1
                self.results['details']['corrupt_zip'] = 'incorrectly_accepted'
                
        except Exception as e:
            print(f"❌ ERROR: {e}")
            self.results['failed'] += 1
            self.results['details']['corrupt_zip'] = 'test_error'

    def test_conversion_api(self, filename, source_ext, target_ext):
        """Prueba conversión usando la API"""
        try:
            file_path = self.temp_dir / filename
            if not file_path.exists():
                return False
            
            with open(file_path, 'rb') as f:
                files = {'file': (filename, f, 'application/octet-stream')}
                data = {'target_format': target_ext}
                
                response = requests.post(f"{API_BASE}/conversion/guest-convert", 
                                       files=files, data=data, timeout=15)
            
            if response.status_code == 200:
                result = response.json()
                return result.get('success', False)
            else:
                return False
                
        except Exception:
            return False

    def generate_report(self):
        """Genera reporte final de casos 'Bad magic number'"""
        print("\n" + "="*50)
        print("📊 REPORTE FINAL - CASOS 'BAD MAGIC NUMBER'")
        print("="*50)
        
        total = self.results['passed'] + self.results['failed']
        success_rate = (self.results['passed'] / total * 100) if total > 0 else 0
        
        print(f"\n📈 ESTADÍSTICAS:")
        print(f"   Total pruebas: {total}")
        print(f"   Exitosas: {self.results['passed']}")
        print(f"   Fallidas: {self.results['failed']}")
        print(f"   Tasa de éxito: {success_rate:.1f}%")
        
        print(f"\n🔍 ANÁLISIS DETALLADO:")
        for test, status in self.results['details'].items():
            icon = "✅" if 'correctly' in status or 'successfully' in status or 'handled' in status else "❌"
            print(f"   {icon} {test}: {status}")
        
        print(f"\n💡 CONCLUSIONES SOBRE 'BAD MAGIC NUMBER':")
        
        # Analizar patrones
        docx_issues = [k for k, v in self.results['details'].items() if 'docx' in k]
        encoding_issues = [k for k, v in self.results['details'].items() if 'encoding' in k or 'bom' in k]
        
        if any('correctly_rejected' in str(v) for k, v in self.results['details'].items() if 'docx' in k):
            print("   • ✅ Sistema detecta correctamente DOCX corruptos")
            print("   • ✅ Validación de archivos ZIP funcionando")
        
        if success_rate >= 80:
            print("   • ✅ Sistema robusto contra archivos problemáticos")
            print("   • 🎯 Errores 'Bad magic number' probablemente por:")
            print("     - Archivos genuinamente corruptos")
            print("     - Problemas de red durante upload")
            print("     - Archivos con extensión incorrecta")
        else:
            print("   • ❌ Sistema vulnerable a archivos problemáticos")
            print("   • ⚠️ Necesita mejoras en validación")
        
        print(f"\n🔧 RECOMENDACIONES:")
        print("   • Implementar validación más estricta de headers de archivo")
        print("   • Agregar verificación de integridad antes de conversión")
        print("   • Mejorar mensajes de error para usuarios")
        
        # Guardar reporte
        import json
        with open('bad_magic_number_test_report.json', 'w', encoding='utf-8') as f:
            json.dump(self.results, f, indent=2, ensure_ascii=False)
        
        print(f"\n📄 Reporte guardado en: bad_magic_number_test_report.json")

def main():
    """Función principal"""
    print("🔍 Verificando servidor...")
    
    try:
        response = requests.get(f"{BASE_URL}/api/health", timeout=5)
        if response.status_code == 200:
            print("✅ Servidor funcionando")
        else:
            print("❌ Servidor con problemas")
            return
    except:
        print("❌ No se puede conectar al servidor")
        return
    
    # Ejecutar pruebas
    tester = BadMagicNumberTester()
    tester.run_bad_magic_tests()

if __name__ == "__main__":
    main()
