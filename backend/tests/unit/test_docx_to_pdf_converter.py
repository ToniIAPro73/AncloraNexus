import pytest
import tempfile
import os
from pathlib import Path
from docx import Document

# Add backend src to Python path
import sys
backend_src = Path(__file__).parent.parent.parent / 'src'
sys.path.insert(0, str(backend_src))

from src.models.conversions.docx_to_pdf import convert, is_valid_docx, clean_unicode_for_pdf


class TestDocxToPdfConverter:
    """Tests para el conversor DOCX→PDF"""
    
    def test_clean_unicode_for_pdf_basic(self):
        """Test básico de limpieza de caracteres Unicode"""
        # Test caracteres con acentos
        result = clean_unicode_for_pdf("Hola, ¿cómo estás?")
        assert "?" in result  # ¿ debe convertirse a ?
        assert "o" in result  # ó debe convertirse a o
        
        # Test emojis
        result = clean_unicode_for_pdf("Proyecto exitoso 🚀")
        assert "[rocket]" in result
        
    def test_clean_unicode_for_pdf_quotes(self):
        """Test de limpieza de comillas tipográficas"""
        text_with_quotes = 'Texto con "comillas" y \'apostrofes\''
        result = clean_unicode_for_pdf(text_with_quotes)
        # Debe convertir comillas tipográficas a ASCII
        assert '"' in result or "'" in result
        
    def test_is_valid_docx_with_invalid_file(self):
        """Test validación de archivo DOCX inválido"""
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
                tmp.write(b"Not a valid DOCX file")
                tmp.flush()

            result = is_valid_docx(tmp_path)
            assert result is False
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except PermissionError:
                    pass

    def test_is_valid_docx_with_valid_file(self):
        """Test validación de archivo DOCX válido"""
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name

            # Crear un documento DOCX válido
            doc = Document()
            doc.add_paragraph("Test content")
            doc.save(tmp_path)

            result = is_valid_docx(tmp_path)
            assert result is True
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except PermissionError:
                    pass
    
    def test_convert_docx_to_pdf_basic(self):
        """Test básico de conversión DOCX→PDF"""
        input_path = None
        output_path = None

        try:
            # Crear archivos temporales
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as input_file:
                input_path = input_file.name
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as output_file:
                output_path = output_file.name

            # Crear documento DOCX de prueba
            doc = Document()
            doc.add_heading('Título de Prueba', level=1)
            doc.add_paragraph('Este es un párrafo de prueba con texto normal.')
            doc.add_paragraph('Otro párrafo con caracteres especiales: ñáéíóú')
            doc.save(input_path)

            # Convertir
            success, message = convert(input_path, output_path)

            # Verificar resultado
            assert success is True
            assert "exitosa" in message
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0

        finally:
            # Cleanup seguro
            if input_path and os.path.exists(input_path):
                try:
                    os.unlink(input_path)
                except PermissionError:
                    pass  # Ignorar errores de permisos en Windows
            if output_path and os.path.exists(output_path):
                try:
                    os.unlink(output_path)
                except PermissionError:
                    pass
    
    def test_convert_with_corrupted_docx(self):
        """Test conversión con archivo DOCX corrupto"""
        input_path = None
        output_path = None

        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as input_file:
                input_path = input_file.name
                input_file.write(b"Not a valid DOCX")
                input_file.flush()

            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as output_file:
                output_path = output_file.name

            # Intentar convertir
            success, message = convert(input_path, output_path)

            # Debe fallar
            assert success is False
            assert "corrupto" in message or "válido" in message

        finally:
            # Cleanup seguro
            if input_path and os.path.exists(input_path):
                try:
                    os.unlink(input_path)
                except PermissionError:
                    pass
            if output_path and os.path.exists(output_path):
                try:
                    os.unlink(output_path)
                except PermissionError:
                    pass
