import os, shutil, tempfile, uuid
from fpdf import FPDF
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
import markdown
import re
from html import unescape
import logging

# Importar motor Pandoc para conversión de alta calidad
try:
    from .pandoc_engine import pandoc_engine, PANDOC_AVAILABLE
except ImportError:
    PANDOC_AVAILABLE = False
    logging.warning("Pandoc engine no disponible para MD→DOCX")

CONVERSION = ('md', 'docx')

def convert(input_path, output_path):
    """Convierte Markdown a DOCX usando Pandoc (preserva Unicode) con fallback manual"""

    # Método 1: Pandoc (RECOMENDADO - preserva formato y Unicode perfectamente)
    if PANDOC_AVAILABLE:
        try:
            success, message = pandoc_engine.convert_with_pandoc(
                input_path, output_path, 'md', 'docx',
                extra_args=[
                    '--wrap=auto',
                    '--columns=80'
                ]
            )
            if success:
                return True, f"Conversión MD→DOCX exitosa con Pandoc (formato preservado) - {message}"
        except Exception as e:
            logging.warning(f"Pandoc falló para MD→DOCX: {e}, usando fallback manual")

    # Método 2: Fallback manual (procesamiento línea por línea)
    return convert_with_manual_processing(input_path, output_path)

def convert_with_manual_processing(input_path, output_path):
    """Fallback usando procesamiento manual de Markdown"""
    try:
        # Leer contenido markdown
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Crear documento Word
        doc = Document()

        # Procesar contenido Markdown línea por línea
        process_markdown_to_docx(md_content, doc)

        # Guardar documento
        doc.save(output_path)

        return True, "Conversión MD→DOCX exitosa con procesamiento manual (Unicode preservado)"
    except Exception as e:
        return False, f"Error en conversión MD→DOCX manual: {str(e)}"

def process_markdown_to_docx(md_content, doc):
    """Procesa contenido Markdown y lo convierte a formato DOCX"""
    lines = md_content.split('\n')
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        line = line.rstrip()
        
        # Manejar bloques de código
        if line.startswith('```'):
            if in_code_block:
                # Finalizar bloque de código
                if code_block_content:
                    add_code_block(doc, '\n'.join(code_block_content))
                code_block_content = []
                in_code_block = False
            else:
                # Iniciar bloque de código
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Líneas vacías
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Títulos
        if line.startswith('#'):
            add_heading(doc, line)
        
        # Listas
        elif line.strip().startswith(('- ', '* ', '+ ')):
            add_bullet_point(doc, line)
        
        elif re.match(r'^\d+\.\s', line.strip()):
            add_numbered_point(doc, line)
        
        # Citas
        elif line.strip().startswith('>'):
            add_quote(doc, line)
        
        # Tablas (formato simple)
        elif '|' in line and line.count('|') >= 2:
            add_table_row(doc, line)
        
        # Enlaces
        elif '[' in line and '](' in line:
            add_paragraph_with_links(doc, line)
        
        # Párrafo normal con formato
        else:
            add_formatted_paragraph(doc, line)

def add_heading(doc, line):
    """Agrega un título al documento"""
    # Contar # para determinar el nivel
    level = 0
    for char in line:
        if char == '#':
            level += 1
        else:
            break
    
    # Limitar nivel a 3
    level = min(level, 3)
    
    # Extraer texto del título
    title_text = line.lstrip('#').strip()
    
    # Agregar título
    heading = doc.add_heading(title_text, level=level)

def add_bullet_point(doc, line):
    """Agrega un punto de lista"""
    # Extraer texto sin el marcador
    text = re.sub(r'^[\s]*[-*+]\s*', '', line)
    
    # Agregar párrafo con formato de lista
    p = doc.add_paragraph(text, style='List Bullet')

def add_numbered_point(doc, line):
    """Agrega un punto numerado"""
    # Extraer texto sin el número
    text = re.sub(r'^\d+\.\s*', '', line.strip())
    
    # Agregar párrafo con formato de lista numerada
    p = doc.add_paragraph(text, style='List Number')

def add_quote(doc, line):
    """Agrega una cita"""
    # Extraer texto de la cita
    quote_text = line.lstrip('>').strip()
    
    # Agregar párrafo con estilo de cita
    p = doc.add_paragraph(quote_text)
    p.style = 'Quote'

def add_code_block(doc, code_content):
    """Agrega un bloque de código"""
    # Agregar párrafo con estilo de código
    p = doc.add_paragraph(code_content)
    # Aplicar formato monospace (si está disponible)
    for run in p.runs:
        run.font.name = 'Courier New'

def add_table_row(doc, line):
    """Agrega una fila de tabla (implementación básica)"""
    # Esta es una implementación simple - en un caso real necesitaríamos
    # manejar tablas completas, no solo filas individuales
    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
    
    if cells:
        # Crear tabla si no existe o agregar fila
        # Por simplicidad, convertimos a párrafo con formato especial
        table_text = ' | '.join(cells)
        p = doc.add_paragraph(table_text)
        for run in p.runs:
            run.font.name = 'Courier New'

def add_paragraph_with_links(doc, line):
    """Agrega párrafo con enlaces"""
    # Implementación básica - extraer enlaces y texto
    import re
    
    # Buscar patrones de enlaces [texto](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    p = doc.add_paragraph()
    last_end = 0
    
    for match in re.finditer(link_pattern, line):
        # Agregar texto antes del enlace
        if match.start() > last_end:
            p.add_run(line[last_end:match.start()])
        
        # Agregar enlace (como texto por ahora)
        link_text = match.group(1)
        link_url = match.group(2)
        p.add_run(f"{link_text} ({link_url})")
        
        last_end = match.end()
    
    # Agregar texto restante
    if last_end < len(line):
        p.add_run(line[last_end:])

def add_formatted_paragraph(doc, line):
    """Agrega párrafo con formato (negrita, cursiva, etc.)"""
    p = doc.add_paragraph()
    
    # Procesar formato inline
    process_inline_formatting(p, line)

def process_inline_formatting(paragraph, text):
    """Procesa formato inline como **negrita** y *cursiva*"""
    import re
    
    # Patrones para formato
    patterns = [
        (r'\*\*([^*]+)\*\*', 'bold'),      # **negrita**
        (r'\*([^*]+)\*', 'italic'),        # *cursiva*
        (r'`([^`]+)`', 'code'),            # `código`
    ]
    
    current_text = text
    segments = []
    
    # Encontrar todos los segmentos con formato
    for pattern, format_type in patterns:
        matches = list(re.finditer(pattern, current_text))
        for match in reversed(matches):  # Procesar de atrás hacia adelante
            start, end = match.span()
            
            # Agregar segmento con formato
            segments.insert(0, {
                'text': match.group(1),
                'format': format_type,
                'start': start,
                'end': end
            })
            
            # Reemplazar en el texto actual
            current_text = current_text[:start] + f"__PLACEHOLDER_{len(segments)-1}__" + current_text[end:]
    
    # Procesar texto con placeholders
    parts = current_text.split('__PLACEHOLDER_')
    
    # Agregar primera parte (texto normal)
    if parts[0]:
        paragraph.add_run(parts[0])
    
    # Agregar segmentos con formato
    for i, segment in enumerate(segments):
        run = paragraph.add_run(segment['text'])
        
        if segment['format'] == 'bold':
            run.bold = True
        elif segment['format'] == 'italic':
            run.italic = True
        elif segment['format'] == 'code':
            run.font.name = 'Courier New'
        
        # Agregar texto después del placeholder si existe
        placeholder_parts = parts[i + 1].split('_', 1)
        if len(placeholder_parts) > 1 and placeholder_parts[1]:
            paragraph.add_run(placeholder_parts[1])
    
    # Si no hay formato especial, agregar texto normal
    if not segments and text:
        paragraph.add_run(text)
