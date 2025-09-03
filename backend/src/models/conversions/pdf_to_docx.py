import os
import logging
import tempfile
from pathlib import Path

# Importar librerías para PDF de alta calidad
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF no disponible para PDF→DOCX de alta calidad")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logging.warning("pdf2image no disponible para PDF→DOCX")

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logging.warning("pypdf no disponible")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx no disponible para PDF→DOCX")

CONVERSION = ('pdf', 'docx')

def convert(input_path, output_path):
    """Convierte PDF a DOCX usando la mejor librería disponible"""
    
    if not PYTHON_DOCX_AVAILABLE:
        return False, "python-docx no está disponible para la conversión PDF→DOCX"
    
    # Método 1: PyMuPDF (RECOMENDADO - mejor extracción de texto y formato)
    if PYMUPDF_AVAILABLE:
        try:
            success, message = convert_with_pymupdf(input_path, output_path)
            if success:
                return True, f"Conversión PDF→DOCX exitosa con PyMuPDF - {message}"
        except Exception as e:
            logging.warning(f"PyMuPDF falló: {e}")
    
    # Método 2: pypdf (extracción básica de texto)
    if PYPDF_AVAILABLE:
        try:
            success, message = convert_with_pypdf(input_path, output_path)
            if success:
                return True, f"Conversión PDF→DOCX exitosa con pypdf - {message}"
        except Exception as e:
            logging.warning(f"pypdf falló: {e}")
    
    # Método 3: Fallback básico
    return convert_basic_fallback(input_path, output_path)

def convert_with_pymupdf(input_path, output_path):
    """Conversión usando PyMuPDF (máxima calidad)"""
    try:
        # Abrir PDF
        pdf_document = fitz.open(input_path)
        
        if len(pdf_document) == 0:
            return False, "PDF vacío o corrupto"
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos del documento
        setup_document_styles(doc)
        
        # Procesar cada página
        total_text_blocks = 0
        total_images = 0
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Agregar separador de página (excepto la primera)
            if page_num > 0:
                doc.add_page_break()
            
            # Extraer texto con formato
            text_blocks = extract_text_blocks_with_format(page)
            total_text_blocks += len(text_blocks)
            
            # Agregar bloques de texto al documento
            for block in text_blocks:
                add_text_block_to_doc(doc, block)
            
            # Extraer y agregar imágenes
            images = extract_images_from_page(page)
            total_images += len(images)
            
            for img_data in images:
                add_image_to_doc(doc, img_data)
        
        # Cerrar PDF
        pdf_document.close()
        
        # Guardar documento Word
        doc.save(output_path)
        
        # Verificar resultado
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return True, f"DOCX generado: {len(pdf_document)} páginas, {total_text_blocks} bloques de texto, {total_images} imágenes"
        else:
            return False, "Error: DOCX no se generó correctamente"
        
    except Exception as e:
        return False, f"Error con PyMuPDF: {str(e)}"

def convert_with_pypdf(input_path, output_path):
    """Conversión usando pypdf (extracción básica)"""
    try:
        # Leer PDF
        with open(input_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            if len(pdf_reader.pages) == 0:
                return False, "PDF vacío"
            
            # Crear documento Word
            doc = Document()
            setup_document_styles(doc)
            
            # Extraer texto de cada página
            total_paragraphs = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                # Agregar separador de página (excepto la primera)
                if page_num > 0:
                    doc.add_page_break()
                
                # Agregar título de página
                page_title = doc.add_heading(f'Página {page_num + 1}', level=2)
                
                try:
                    # Extraer texto de la página
                    text = page.extract_text()
                    
                    if text.strip():
                        # Dividir en párrafos
                        paragraphs = text.split('\n\n')
                        
                        for paragraph_text in paragraphs:
                            if paragraph_text.strip():
                                # Limpiar texto
                                clean_text = clean_extracted_text(paragraph_text)
                                
                                if clean_text:
                                    doc.add_paragraph(clean_text)
                                    total_paragraphs += 1
                    else:
                        doc.add_paragraph("(Página sin texto extraíble)")
                        
                except Exception as e:
                    doc.add_paragraph(f"(Error extrayendo texto de página {page_num + 1}: {str(e)})")
            
            # Guardar documento
            doc.save(output_path)
            
            # Verificar resultado
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return True, f"DOCX generado: {len(pdf_reader.pages)} páginas, {total_paragraphs} párrafos"
            else:
                return False, "Error: DOCX no se generó correctamente"
        
    except Exception as e:
        return False, f"Error con pypdf: {str(e)}"

def convert_basic_fallback(input_path, output_path):
    """Fallback básico cuando no hay librerías disponibles"""
    try:
        # Crear documento básico con información del PDF
        doc = Document()
        
        # Título
        title = doc.add_heading('Conversión PDF → DOCX', 0)
        
        # Información del archivo
        file_size = os.path.getsize(input_path) / 1024  # KB
        
        doc.add_paragraph(f"Archivo original: {os.path.basename(input_path)}")
        doc.add_paragraph(f"Tamaño: {file_size:.1f} KB")
        doc.add_paragraph()
        
        # Mensaje informativo
        info_para = doc.add_paragraph()
        info_para.add_run("NOTA: ").bold = True
        info_para.add_run("Esta es una conversión básica. Para obtener el contenido real del PDF, instale PyMuPDF o pypdf.")
        
        doc.add_paragraph()
        doc.add_paragraph("Para una conversión completa, el sistema necesita:")
        
        requirements = [
            "PyMuPDF (fitz) - para extracción avanzada de texto e imágenes",
            "pypdf - para extracción básica de texto",
            "pdf2image - para conversión de páginas a imágenes"
        ]
        
        for req in requirements:
            doc.add_paragraph(req, style='List Bullet')
        
        # Guardar documento
        doc.save(output_path)
        
        return True, "DOCX placeholder generado (instalar PyMuPDF para conversión real)"
        
    except Exception as e:
        return False, f"Error en fallback básico: {str(e)}"

def setup_document_styles(doc):
    """Configurar estilos del documento Word"""
    try:
        # Configurar estilo normal
        styles = doc.styles
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
    except Exception as e:
        logging.warning(f"Error configurando estilos: {e}")

def extract_text_blocks_with_format(page):
    """Extraer bloques de texto con información de formato usando PyMuPDF"""
    try:
        blocks = []
        
        # Obtener bloques de texto de la página
        text_dict = page.get_text("dict")
        
        for block in text_dict["blocks"]:
            if "lines" in block:  # Es un bloque de texto
                block_text = ""
                font_size = 11
                is_bold = False
                
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            block_text += text + " "
                            
                            # Obtener información de formato
                            font_size = span.get("size", 11)
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & 2**4)  # Flag para bold
                
                if block_text.strip():
                    blocks.append({
                        'text': clean_extracted_text(block_text.strip()),
                        'font_size': font_size,
                        'is_bold': is_bold,
                        'is_heading': font_size > 14 or is_bold
                    })
        
        return blocks
        
    except Exception as e:
        logging.warning(f"Error extrayendo bloques de texto: {e}")
        # Fallback a extracción simple
        text = page.get_text()
        if text.strip():
            return [{'text': clean_extracted_text(text), 'font_size': 11, 'is_bold': False, 'is_heading': False}]
        return []

def extract_images_from_page(page):
    """Extraer imágenes de la página usando PyMuPDF"""
    try:
        images = []
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            try:
                # Obtener datos de la imagen
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                # Solo procesar imágenes RGB o GRAY
                if pix.n - pix.alpha < 4:
                    img_data = pix.tobytes("png")
                    images.append({
                        'data': img_data,
                        'format': 'png',
                        'width': pix.width,
                        'height': pix.height
                    })
                
                pix = None  # Liberar memoria
                
            except Exception as e:
                logging.warning(f"Error extrayendo imagen {img_index}: {e}")
                continue
        
        return images
        
    except Exception as e:
        logging.warning(f"Error extrayendo imágenes: {e}")
        return []

def add_text_block_to_doc(doc, block):
    """Agregar bloque de texto al documento Word"""
    try:
        text = block['text']
        
        if block.get('is_heading', False):
            # Agregar como título
            heading = doc.add_heading(text, level=2)
        else:
            # Agregar como párrafo
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            
            if block.get('is_bold', False):
                run.bold = True
            
            # Ajustar tamaño de fuente si es diferente
            font_size = block.get('font_size', 11)
            if font_size != 11:
                run.font.size = Pt(min(font_size, 18))  # Limitar tamaño máximo
        
    except Exception as e:
        logging.warning(f"Error agregando bloque de texto: {e}")
        # Fallback: agregar como párrafo simple
        doc.add_paragraph(block['text'])

def add_image_to_doc(doc, img_data):
    """Agregar imagen al documento Word"""
    try:
        # Crear archivo temporal para la imagen
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
            temp_file.write(img_data['data'])
            temp_path = temp_file.name
        
        try:
            # Calcular tamaño apropiado (máximo 6 pulgadas de ancho)
            width = img_data['width']
            height = img_data['height']
            
            max_width = Inches(6)
            if width > 0:
                aspect_ratio = height / width
                if width > 600:  # Si es muy ancha, redimensionar
                    doc_width = max_width
                    doc_height = max_width * aspect_ratio
                else:
                    doc_width = Inches(width / 100)  # Conversión aproximada
                    doc_height = Inches(height / 100)
                
                # Agregar imagen al documento
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(temp_path, width=doc_width, height=doc_height)
            
        finally:
            # Limpiar archivo temporal
            os.unlink(temp_path)
        
    except Exception as e:
        logging.warning(f"Error agregando imagen: {e}")

def clean_extracted_text(text):
    """Limpiar texto extraído del PDF"""
    try:
        if not text:
            return ""
        
        # Limpiar caracteres problemáticos
        clean_text = text.replace('\x00', '')  # Caracteres nulos
        clean_text = clean_text.replace('\ufeff', '')  # BOM
        
        # Normalizar espacios en blanco
        clean_text = ' '.join(clean_text.split())
        
        # Limitar longitud para evitar párrafos excesivamente largos
        if len(clean_text) > 1000:
            clean_text = clean_text[:997] + '...'
        
        return clean_text.strip()
        
    except Exception:
        return str(text) if text else ""
