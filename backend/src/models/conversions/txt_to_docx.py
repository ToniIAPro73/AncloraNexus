import os, shutil, tempfile, uuid
from fpdf import FPDF
from PIL import Image, ImageDraw
from docx import Document
from pypdf import PdfReader

CONVERSION = ('txt', 'docx')

def convert(input_path, output_path):
    """Convierte TXT a DOCX con mejor manejo de errores"""
    try:
        # Crear documento Word
        doc = Document()
        
        # Leer contenido del archivo TXT
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Procesar contenido línea por línea
        lines = content.split('\n')
        
        # Agregar título si la primera línea parece ser un título
        if lines and len(lines[0].strip()) < 60 and not lines[0].strip().endswith('.'):
            doc.add_heading(lines[0].strip(), level=1)
            lines = lines[1:]
        
        # Agregar párrafos
        for line in lines:
            line = line.strip()
            if line:  # Solo agregar líneas no vacías
                doc.add_paragraph(line)
            else:
                # Agregar espacio para líneas vacías
                doc.add_paragraph()
        
        # Asegurar que el directorio de salida existe
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Guardar documento
        doc.save(output_path)
        
        # Verificar que el archivo se guardó correctamente
        if not os.path.exists(output_path):
            return False, "Error: El archivo DOCX no se guardó correctamente"
        
        # Verificar que el archivo tiene contenido
        if os.path.getsize(output_path) == 0:
            return False, "Error: El archivo DOCX está vacío"
        
        return True, "Conversión TXT→DOCX exitosa"
    except Exception as e:
        return False, f"Error en conversión TXT→DOCX: {str(e)}"
