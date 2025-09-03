import os
import logging
import re
import tempfile
from pathlib import Path

# Importar librerías necesarias
try:
    from bs4 import BeautifulSoup, NavigableString
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    logging.warning("BeautifulSoup no disponible para HTML→DOCX")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx no disponible para HTML→DOCX")

try:
    import requests
    from urllib.parse import urljoin, urlparse
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False
    logging.warning("requests no disponible para descargar imágenes en HTML→DOCX")

CONVERSION = ('html', 'docx')

def convert(input_path, output_path):
    """Convierte HTML a DOCX preservando estructura y formato"""
    
    if not PYTHON_DOCX_AVAILABLE:
        return False, "python-docx no está disponible para la conversión HTML→DOCX"
    
    if not BS4_AVAILABLE:
        return False, "BeautifulSoup no está disponible para la conversión HTML→DOCX"
    
    try:
        # Leer archivo HTML
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        if not html_content.strip():
            return False, "Archivo HTML vacío"
        
        # Parsear HTML con BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos del documento
        setup_document_styles(doc)
        
        # Extraer título del documento
        title = extract_document_title(soup)
        if title:
            doc.add_heading(title, 0)
        
        # Procesar el contenido del body
        body = soup.find('body')
        if body:
            process_html_element(doc, body, input_path)
        else:
            # Si no hay body, procesar todo el contenido
            process_html_element(doc, soup, input_path)
        
        # Guardar documento
        doc.save(output_path)
        
        # Verificar resultado
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            # Contar elementos procesados
            stats = count_html_elements(soup)
            return True, f"DOCX generado: {stats}"
        else:
            return False, "Error: DOCX no se generó correctamente"
        
    except Exception as e:
        return False, f"Error en conversión HTML→DOCX: {str(e)}"

def setup_document_styles(doc):
    """Configurar estilos del documento Word"""
    try:
        # Configurar estilo normal
        styles = doc.styles
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Crear estilo para código
        try:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
        except:
            pass  # El estilo ya existe
        
    except Exception as e:
        logging.warning(f"Error configurando estilos: {e}")

def extract_document_title(soup):
    """Extraer título del documento HTML"""
    try:
        # Buscar en orden de prioridad
        title_sources = [
            soup.find('title'),
            soup.find('h1'),
            soup.find('h2'),
            soup.find('meta', attrs={'name': 'title'}),
            soup.find('meta', attrs={'property': 'og:title'})
        ]
        
        for source in title_sources:
            if source:
                if source.name == 'meta':
                    title = source.get('content', '').strip()
                else:
                    title = source.get_text().strip()
                
                if title and len(title) < 200:  # Título razonable
                    return clean_text(title)
        
        return None
        
    except Exception:
        return None

def process_html_element(doc, element, base_path):
    """Procesar elemento HTML y agregarlo al documento"""
    try:
        if isinstance(element, NavigableString):
            # Es texto directo
            text = clean_text(str(element))
            if text:
                doc.add_paragraph(text)
            return
        
        tag_name = element.name.lower() if element.name else ''
        
        # Procesar según el tipo de elemento
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            process_heading(doc, element, tag_name)
        
        elif tag_name == 'p':
            process_paragraph(doc, element)
        
        elif tag_name in ['ul', 'ol']:
            process_list(doc, element)
        
        elif tag_name == 'table':
            process_table(doc, element)
        
        elif tag_name == 'img':
            process_image(doc, element, base_path)
        
        elif tag_name in ['pre', 'code']:
            process_code(doc, element)
        
        elif tag_name == 'blockquote':
            process_blockquote(doc, element)
        
        elif tag_name == 'hr':
            add_horizontal_rule(doc)
        
        elif tag_name in ['div', 'section', 'article', 'main']:
            # Contenedores - procesar hijos
            for child in element.children:
                process_html_element(doc, child, base_path)
        
        elif tag_name in ['br']:
            doc.add_paragraph()  # Salto de línea
        
        else:
            # Para otros elementos, procesar el contenido
            text = element.get_text() if hasattr(element, 'get_text') else str(element)
            text = clean_text(text)
            if text:
                doc.add_paragraph(text)
        
    except Exception as e:
        logging.warning(f"Error procesando elemento {tag_name}: {e}")

def process_heading(doc, element, tag_name):
    """Procesar encabezados HTML"""
    try:
        text = clean_text(element.get_text())
        if text:
            level = int(tag_name[1])  # h1 -> 1, h2 -> 2, etc.
            doc.add_heading(text, level=min(level, 9))  # Word soporta hasta nivel 9
    except Exception as e:
        logging.warning(f"Error procesando encabezado: {e}")

def process_paragraph(doc, element):
    """Procesar párrafos HTML con formato"""
    try:
        paragraph = doc.add_paragraph()
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = clean_text(str(child))
                if text:
                    paragraph.add_run(text)
            else:
                process_inline_element(paragraph, child)
    
    except Exception as e:
        logging.warning(f"Error procesando párrafo: {e}")

def process_inline_element(paragraph, element):
    """Procesar elementos inline (strong, em, etc.)"""
    try:
        tag_name = element.name.lower() if element.name else ''
        text = clean_text(element.get_text())
        
        if not text:
            return
        
        run = paragraph.add_run(text)
        
        # Aplicar formato según el tag
        if tag_name in ['strong', 'b']:
            run.bold = True
        elif tag_name in ['em', 'i']:
            run.italic = True
        elif tag_name == 'u':
            run.underline = True
        elif tag_name in ['code', 'tt']:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif tag_name == 'a':
            # Enlaces - agregar como texto con formato especial
            href = element.get('href', '')
            if href:
                run.font.color.rgb = RGBColor(0, 0, 255)  # Azul
                run.underline = True
    
    except Exception as e:
        logging.warning(f"Error procesando elemento inline: {e}")

def process_list(doc, element):
    """Procesar listas HTML"""
    try:
        is_ordered = element.name.lower() == 'ol'
        
        for li in element.find_all('li', recursive=False):
            text = clean_text(li.get_text())
            if text:
                if is_ordered:
                    doc.add_paragraph(text, style='List Number')
                else:
                    doc.add_paragraph(text, style='List Bullet')
    
    except Exception as e:
        logging.warning(f"Error procesando lista: {e}")

def process_table(doc, element):
    """Procesar tablas HTML"""
    try:
        rows = element.find_all('tr')
        if not rows:
            return
        
        # Determinar número de columnas
        max_cols = 0
        for row in rows:
            cols = len(row.find_all(['td', 'th']))
            max_cols = max(max_cols, cols)
        
        if max_cols == 0:
            return
        
        # Crear tabla en Word
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # Llenar tabla
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    cell_text = clean_text(cell.get_text())
                    table.cell(row_idx, col_idx).text = cell_text
                    
                    # Formatear encabezados
                    if cell.name == 'th':
                        table.cell(row_idx, col_idx).paragraphs[0].runs[0].bold = True
    
    except Exception as e:
        logging.warning(f"Error procesando tabla: {e}")

def process_image(doc, element, base_path):
    """Procesar imágenes HTML"""
    try:
        src = element.get('src', '')
        alt = element.get('alt', '')
        
        if not src:
            return
        
        # Intentar descargar/acceder a la imagen
        image_path = None
        
        # Si es una ruta relativa, intentar resolverla
        if not src.startswith(('http://', 'https://', 'data:')):
            base_dir = os.path.dirname(base_path)
            potential_path = os.path.join(base_dir, src)
            if os.path.exists(potential_path):
                image_path = potential_path
        
        # Si es una URL y requests está disponible, intentar descargar
        elif src.startswith(('http://', 'https://')) and REQUESTS_AVAILABLE:
            try:
                response = requests.get(src, timeout=10)
                if response.status_code == 200:
                    # Crear archivo temporal
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                        temp_file.write(response.content)
                        image_path = temp_file.name
            except:
                pass
        
        # Agregar imagen si se pudo obtener
        if image_path and os.path.exists(image_path):
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(4))  # Tamaño fijo
                
                # Agregar texto alternativo como descripción
                if alt:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(alt)
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
                
                # Limpiar archivo temporal si se creó
                if image_path.startswith(tempfile.gettempdir()):
                    os.unlink(image_path)
                    
            except Exception as e:
                logging.warning(f"Error insertando imagen: {e}")
        else:
            # Si no se pudo cargar la imagen, agregar texto alternativo
            if alt or src:
                doc.add_paragraph(f"[Imagen: {alt or src}]")
    
    except Exception as e:
        logging.warning(f"Error procesando imagen: {e}")

def process_code(doc, element):
    """Procesar bloques de código"""
    try:
        code_text = element.get_text()
        if code_text:
            # Preservar espacios y saltos de línea
            lines = code_text.split('\n')
            for line in lines:
                para = doc.add_paragraph(line)
                para.style = 'Code' if 'Code' in doc.styles else 'Normal'
                if 'Code' not in doc.styles:
                    # Aplicar formato de código manualmente
                    for run in para.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
    
    except Exception as e:
        logging.warning(f"Error procesando código: {e}")

def process_blockquote(doc, element):
    """Procesar citas"""
    try:
        text = clean_text(element.get_text())
        if text:
            para = doc.add_paragraph(text)
            para.style = 'Quote'
    
    except Exception as e:
        logging.warning(f"Error procesando cita: {e}")

def add_horizontal_rule(doc):
    """Agregar línea horizontal"""
    try:
        para = doc.add_paragraph()
        para.add_run("_" * 50)  # Línea de guiones bajos
    except Exception as e:
        logging.warning(f"Error agregando línea horizontal: {e}")

def clean_text(text):
    """Limpiar texto HTML"""
    try:
        if not text:
            return ""
        
        # Convertir a string si no lo es
        text = str(text)
        
        # Limpiar espacios en blanco excesivos
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # Decodificar entidades HTML básicas
        html_entities = {
            '&amp;': '&',
            '&lt;': '<',
            '&gt;': '>',
            '&quot;': '"',
            '&apos;': "'",
            '&nbsp;': ' '
        }
        
        for entity, char in html_entities.items():
            text = text.replace(entity, char)
        
        return text
        
    except Exception:
        return str(text) if text else ""

def count_html_elements(soup):
    """Contar elementos HTML procesados"""
    try:
        stats = []
        
        # Contar diferentes tipos de elementos
        headings = len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']))
        paragraphs = len(soup.find_all('p'))
        lists = len(soup.find_all(['ul', 'ol']))
        tables = len(soup.find_all('table'))
        images = len(soup.find_all('img'))
        links = len(soup.find_all('a'))
        
        if headings > 0:
            stats.append(f"{headings} encabezados")
        if paragraphs > 0:
            stats.append(f"{paragraphs} párrafos")
        if lists > 0:
            stats.append(f"{lists} listas")
        if tables > 0:
            stats.append(f"{tables} tablas")
        if images > 0:
            stats.append(f"{images} imágenes")
        if links > 0:
            stats.append(f"{links} enlaces")
        
        return ', '.join(stats) if stats else "contenido HTML procesado"
        
    except Exception:
        return "contenido HTML procesado"
