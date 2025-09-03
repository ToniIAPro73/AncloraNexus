import os
import logging
import tempfile
from pathlib import Path

# Importar librerías necesarias
try:
    from PIL import Image, ExifTags
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logging.warning("PIL no disponible para JPG→DOCX")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx no disponible para JPG→DOCX")

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logging.warning("pytesseract no disponible para OCR en JPG→DOCX")

CONVERSION = ('jpg', 'docx')

def convert(input_path, output_path):
    """Convierte JPG a DOCX insertando la imagen con información adicional"""
    
    if not PYTHON_DOCX_AVAILABLE:
        return False, "python-docx no está disponible para la conversión JPG→DOCX"
    
    if not PIL_AVAILABLE:
        return False, "PIL no está disponible para la conversión JPG→DOCX"
    
    try:
        # Validar archivo JPG
        is_valid, validation_message = validate_jpg_file(input_path)
        if not is_valid:
            return False, f"Archivo JPG inválido: {validation_message}"
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos del documento
        setup_document_styles(doc)
        
        # Agregar título
        title = doc.add_heading('Imagen JPG → Documento Word', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Obtener información de la imagen
        img_info = get_image_info(input_path)
        
        # Agregar información de la imagen
        add_image_info_section(doc, input_path, img_info)
        
        # Agregar la imagen al documento
        success, message = add_image_to_document(doc, input_path, img_info)
        if not success:
            return False, f"Error agregando imagen: {message}"
        
        # Intentar OCR si está disponible
        if TESSERACT_AVAILABLE:
            ocr_text = extract_text_with_ocr(input_path)
            if ocr_text:
                add_ocr_section(doc, ocr_text)
        
        # Agregar metadatos EXIF si están disponibles
        if img_info.get('exif_data'):
            add_exif_section(doc, img_info['exif_data'])
        
        # Guardar documento
        doc.save(output_path)
        
        # Verificar resultado
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return True, f"DOCX generado: imagen {img_info['width']}x{img_info['height']}px insertada con información completa"
        else:
            return False, "Error: DOCX no se generó correctamente"
        
    except Exception as e:
        return False, f"Error en conversión JPG→DOCX: {str(e)}"

def validate_jpg_file(file_path):
    """Validar que el archivo es un JPG válido"""
    try:
        # Verificar extensión
        if not file_path.lower().endswith(('.jpg', '.jpeg')):
            return False, "Archivo no tiene extensión JPG/JPEG"
        
        # Verificar que existe y no está vacío
        if not os.path.exists(file_path):
            return False, "Archivo no existe"
        
        if os.path.getsize(file_path) == 0:
            return False, "Archivo está vacío"
        
        # Verificar magic number de JPEG
        with open(file_path, 'rb') as f:
            header = f.read(4)
            
            if len(header) < 4:
                return False, "Archivo demasiado pequeño para ser JPG válido"
            
            # JPEG comienza con FF D8 FF
            if header[:2] != b'\xff\xd8':
                return False, "No es un archivo JPEG válido"
        
        # Intentar abrir con PIL para validación adicional
        if PIL_AVAILABLE:
            try:
                with Image.open(file_path) as img:
                    img.verify()  # Verificar integridad
            except Exception as e:
                return False, f"Imagen corrupta: {str(e)}"
        
        return True, "Archivo JPG válido"
        
    except Exception as e:
        return False, f"Error validando archivo: {str(e)}"

def get_image_info(file_path):
    """Obtener información completa de la imagen"""
    try:
        info = {}
        
        # Información básica del archivo
        info['file_name'] = os.path.basename(file_path)
        info['file_size'] = os.path.getsize(file_path)
        info['file_size_kb'] = info['file_size'] / 1024
        info['file_size_mb'] = info['file_size'] / (1024 * 1024)
        
        # Información de la imagen con PIL
        if PIL_AVAILABLE:
            with Image.open(file_path) as img:
                info['width'] = img.width
                info['height'] = img.height
                info['mode'] = img.mode
                info['format'] = img.format
                
                # Calcular megapíxeles
                info['megapixels'] = (img.width * img.height) / 1000000
                
                # Calcular relación de aspecto
                gcd_val = gcd(img.width, img.height)
                info['aspect_ratio'] = f"{img.width // gcd_val}:{img.height // gcd_val}"
                
                # Obtener datos EXIF si están disponibles
                info['exif_data'] = extract_exif_data(img)
        
        return info
        
    except Exception as e:
        return {'error': str(e)}

def extract_exif_data(img):
    """Extraer datos EXIF de la imagen"""
    try:
        exif_data = {}
        
        if hasattr(img, '_getexif') and img._getexif() is not None:
            exif = img._getexif()
            
            for tag_id, value in exif.items():
                tag = ExifTags.TAGS.get(tag_id, tag_id)
                
                # Filtrar tags más importantes
                important_tags = [
                    'DateTime', 'DateTimeOriginal', 'DateTimeDigitized',
                    'Make', 'Model', 'Software',
                    'ExposureTime', 'FNumber', 'ISO', 'FocalLength',
                    'Flash', 'WhiteBalance', 'ExposureMode',
                    'ImageWidth', 'ImageLength', 'Orientation'
                ]
                
                if tag in important_tags:
                    # Convertir valores complejos a string
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8')
                        except:
                            value = str(value)
                    elif isinstance(value, tuple) and len(value) == 2:
                        # Fracciones (como tiempo de exposición)
                        if value[1] != 0:
                            value = f"{value[0]}/{value[1]}"
                        else:
                            value = str(value[0])
                    
                    exif_data[tag] = str(value)
        
        return exif_data
        
    except Exception as e:
        logging.warning(f"Error extrayendo EXIF: {e}")
        return {}

def setup_document_styles(doc):
    """Configurar estilos del documento Word"""
    try:
        # Configurar estilo normal
        styles = doc.styles
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
    except Exception as e:
        logging.warning(f"Error configurando estilos: {e}")

def add_image_info_section(doc, file_path, img_info):
    """Agregar sección con información de la imagen"""
    try:
        # Título de sección
        doc.add_heading('Información de la Imagen', level=1)
        
        # Información básica
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'
        
        # Datos a mostrar
        info_items = [
            ('Nombre del archivo', img_info.get('file_name', 'N/A')),
            ('Tamaño del archivo', f"{img_info.get('file_size_kb', 0):.1f} KB ({img_info.get('file_size_mb', 0):.2f} MB)"),
            ('Dimensiones', f"{img_info.get('width', 0)} × {img_info.get('height', 0)} píxeles"),
            ('Megapíxeles', f"{img_info.get('megapixels', 0):.1f} MP"),
            ('Relación de aspecto', img_info.get('aspect_ratio', 'N/A')),
            ('Modo de color', img_info.get('mode', 'N/A')),
            ('Formato', img_info.get('format', 'N/A'))
        ]
        
        for label, value in info_items:
            row = info_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            # Formatear primera columna en negrita
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Espacio
        
    except Exception as e:
        logging.warning(f"Error agregando información de imagen: {e}")

def add_image_to_document(doc, file_path, img_info):
    """Agregar la imagen al documento con tamaño apropiado"""
    try:
        # Título de sección
        doc.add_heading('Imagen', level=1)
        
        # Calcular tamaño apropiado para el documento
        original_width = img_info.get('width', 800)
        original_height = img_info.get('height', 600)
        
        # Tamaño máximo en el documento (6.5 pulgadas de ancho máximo)
        max_width = Inches(6.5)
        max_height = Inches(8)
        
        # Calcular escala manteniendo proporción
        width_scale = max_width.inches / (original_width / 96)  # Asumiendo 96 DPI
        height_scale = max_height.inches / (original_height / 96)
        scale = min(width_scale, height_scale, 1.0)  # No agrandar
        
        # Dimensiones finales
        doc_width = Inches((original_width / 96) * scale)
        doc_height = Inches((original_height / 96) * scale)
        
        # Agregar párrafo centrado para la imagen
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar imagen
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(file_path, width=doc_width, height=doc_height)
        
        # Agregar descripción
        doc.add_paragraph()
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(f"Imagen original: {original_width} × {original_height} píxeles")
        caption_run.italic = True
        caption_run.font.size = Pt(9)
        
        return True, f"Imagen insertada: {doc_width.inches:.1f}\" × {doc_height.inches:.1f}\""
        
    except Exception as e:
        return False, f"Error insertando imagen: {str(e)}"

def extract_text_with_ocr(file_path):
    """Extraer texto de la imagen usando OCR"""
    try:
        if not TESSERACT_AVAILABLE:
            return None
        
        # Abrir imagen
        with Image.open(file_path) as img:
            # Convertir a RGB si es necesario
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Aplicar OCR
            text = pytesseract.image_to_string(img, lang='spa+eng')  # Español e inglés
            
            # Limpiar texto
            text = text.strip()
            
            # Solo devolver si hay texto significativo
            if len(text) > 10 and not text.isspace():
                return text
        
        return None
        
    except Exception as e:
        logging.warning(f"Error en OCR: {e}")
        return None

def add_ocr_section(doc, ocr_text):
    """Agregar sección con texto extraído por OCR"""
    try:
        doc.add_heading('Texto Extraído (OCR)', level=1)
        
        # Agregar nota explicativa
        note = doc.add_paragraph()
        note_run = note.add_run("Nota: ")
        note_run.bold = True
        note.add_run("El siguiente texto fue extraído automáticamente de la imagen usando OCR. Puede contener errores.")
        note_run.font.size = Pt(9)
        note_run.italic = True
        
        doc.add_paragraph()
        
        # Agregar texto en un cuadro
        text_para = doc.add_paragraph()
        text_para.style = 'Quote'
        text_para.add_run(ocr_text)
        
    except Exception as e:
        logging.warning(f"Error agregando sección OCR: {e}")

def add_exif_section(doc, exif_data):
    """Agregar sección con datos EXIF"""
    try:
        if not exif_data:
            return
        
        doc.add_heading('Metadatos EXIF', level=1)
        
        # Crear tabla para datos EXIF
        exif_table = doc.add_table(rows=0, cols=2)
        exif_table.style = 'Table Grid'
        
        # Agregar encabezados
        header_row = exif_table.add_row()
        header_row.cells[0].text = 'Campo'
        header_row.cells[1].text = 'Valor'
        
        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Agregar datos EXIF
        for tag, value in exif_data.items():
            row = exif_table.add_row()
            row.cells[0].text = tag
            row.cells[1].text = str(value)
        
    except Exception as e:
        logging.warning(f"Error agregando sección EXIF: {e}")

def gcd(a, b):
    """Calcular máximo común divisor"""
    while b:
        a, b = b, a % b
    return a
