import os, shutil, tempfile, uuid
from fpdf import FPDF
from PIL import Image, ImageDraw
from docx import Document
from pypdf import PdfReader

# Importar motor Pandoc
try:
    from .pandoc_engine import pandoc_engine
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

CONVERSION = ('rtf', 'docx')

def convert(input_path, output_path):
    """Convierte RTF a DOCX usando Pandoc con fallback manual"""
    try:
        # Intentar conversión con Pandoc primero
        if PANDOC_AVAILABLE and pandoc_engine.is_supported_conversion('rtf', 'docx'):
            success, message = pandoc_engine.convert_with_pandoc(
                input_path, output_path, 'rtf', 'docx',
                extra_args=['--wrap=auto', '--columns=80']
            )
            if success:
                return success, message
        
        # Fallback: conversión manual básica
        return convert_rtf_manual(input_path, output_path)
        
    except Exception as e:
        return False, f"Error en conversión RTF→DOCX: {str(e)}"

def convert_rtf_manual(input_path, output_path):
    """Conversión RTF manual como fallback"""
    try:
        # Leer archivo RTF y extraer texto básico
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()
        
        # Extraer texto básico del RTF (muy simplificado)
        text_content = extract_text_from_rtf(rtf_content)
        
        # Crear documento DOCX
        doc = Document()
        
        # Agregar contenido línea por línea
        for line in text_content.split('\n'):
            line = line.strip()
            if line:
                # Detectar títulos básicos
                if line.isupper() and len(line) < 60:
                    doc.add_heading(line, level=1)
                elif line.endswith(':') and len(line) < 50:
                    doc.add_heading(line.rstrip(':'), level=2)
                else:
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph()  # Línea vacía
        
        # Guardar documento
        doc.save(output_path)
        
        # Verificar que se creó correctamente
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            return False, "Error: No se pudo generar el archivo DOCX"
        
        return True, "Conversión RTF→DOCX exitosa (fallback manual)"
        
    except Exception as e:
        return False, f"Error en conversión manual RTF→DOCX: {str(e)}"

def extract_text_from_rtf(rtf_content):
    """Extrae texto básico de contenido RTF"""
    import re
    
    # Remover comandos RTF básicos
    text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)  # Comandos RTF
    text = re.sub(r'[{}]', '', text)  # Llaves RTF
    text = re.sub(r'\\[\'"][a-f0-9]{2}', '', text)  # Caracteres especiales
    
    # Limpiar espacios múltiples
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()
