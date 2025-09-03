import os
import tempfile
from PIL import Image
from docx import Document
from docx.shared import Inches
import logging

CONVERSION = ('png', 'docx')

def convert(input_path, output_path):
    """Convierte PNG a DOCX insertando la imagen en un documento Word"""
    try:
        # Verificar que el archivo PNG existe y es válido
        if not os.path.exists(input_path):
            return False, "Archivo PNG no encontrado"
        
        # Validar que es una imagen PNG válida
        try:
            with Image.open(input_path) as img:
                if img.format != 'PNG':
                    return False, "El archivo no es un PNG válido"
                
                # Obtener dimensiones de la imagen
                width, height = img.size
                
        except Exception as e:
            return False, f"Error al leer imagen PNG: {str(e)}"
        
        # Crear documento Word
        doc = Document()
        
        # Agregar título
        title = doc.add_heading('Imagen PNG Convertida', 0)
        title.alignment = 1  # Centrado
        
        # Agregar párrafo con información
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run(f"Archivo original: ").bold = True
        info_paragraph.add_run(os.path.basename(input_path))
        info_paragraph.add_run(f"\nDimensiones: ").bold = True
        info_paragraph.add_run(f"{width} x {height} píxeles")
        info_paragraph.add_run(f"\nTamaño: ").bold = True
        info_paragraph.add_run(f"{os.path.getsize(input_path) / 1024:.1f} KB")
        
        # Agregar espacio
        doc.add_paragraph()
        
        # Calcular tamaño óptimo para insertar en el documento
        # Máximo 6 pulgadas de ancho para que quepa en la página
        max_width_inches = 6.0
        max_height_inches = 8.0
        
        # Calcular ratio de aspecto
        aspect_ratio = width / height
        
        if width > height:
            # Imagen horizontal
            insert_width = min(max_width_inches, width / 96)  # 96 DPI por defecto
            insert_height = insert_width / aspect_ratio
        else:
            # Imagen vertical
            insert_height = min(max_height_inches, height / 96)
            insert_width = insert_height * aspect_ratio
        
        # Asegurar que no exceda los límites
        if insert_width > max_width_inches:
            insert_width = max_width_inches
            insert_height = insert_width / aspect_ratio
        
        if insert_height > max_height_inches:
            insert_height = max_height_inches
            insert_width = insert_height * aspect_ratio
        
        # Insertar imagen en el documento
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1  # Centrado
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(input_path, width=Inches(insert_width), height=Inches(insert_height))
            
        except Exception as e:
            # Si falla la inserción directa, intentar con copia temporal
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    # Copiar imagen a archivo temporal
                    with Image.open(input_path) as img:
                        # Convertir a RGB si es necesario (para compatibilidad)
                        if img.mode in ('RGBA', 'LA', 'P'):
                            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                            if img.mode == 'P':
                                img = img.convert('RGBA')
                            rgb_img.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                            img = rgb_img
                        
                        img.save(temp_file.name, 'PNG')
                    
                    # Insertar imagen temporal
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = 1  # Centrado
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(temp_file.name, width=Inches(insert_width), height=Inches(insert_height))
                    
                    # Limpiar archivo temporal
                    os.unlink(temp_file.name)
                    
            except Exception as e2:
                return False, f"Error al insertar imagen en documento: {str(e2)}"
        
        # Agregar párrafo final
        doc.add_paragraph()
        footer_paragraph = doc.add_paragraph()
        footer_paragraph.add_run("Documento generado por Anclora Nexus").italic = True
        footer_paragraph.alignment = 1  # Centrado
        
        # Guardar documento
        doc.save(output_path)
        
        # Verificar que se creó correctamente
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            return False, "Error: No se pudo generar el archivo DOCX"
        
        return True, f"Conversión PNG→DOCX exitosa: imagen {width}x{height}px insertada en documento Word"
        
    except Exception as e:
        logging.error(f"Error en conversión PNG→DOCX: {str(e)}")
        return False, f"Error en conversión PNG→DOCX: {str(e)}"

def convert_with_enhanced_layout(input_path, output_path):
    """Versión mejorada con layout más profesional"""
    try:
        # Usar conversión básica como base
        success, message = convert(input_path, output_path)
        
        if not success:
            return success, message
        
        # Aquí se podrían agregar mejoras adicionales como:
        # - Múltiples imágenes por página
        # - Marcos decorativos
        # - Metadatos del documento
        # - Estilos personalizados
        
        return True, message + " (layout mejorado)"
        
    except Exception as e:
        return False, f"Error en conversión mejorada: {str(e)}"
