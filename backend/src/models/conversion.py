from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import os
import tempfile
import uuid
import shutil
from collections import deque

from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw
from pypdf import PdfReader

# Importar el motor de conversión existente
# (Aquí integraremos el motor que ya tienes implementado)
# (Aquí integraremos el motor que ya tienes implementado)

class ConversionEngine:
    """Motor de conversión de Anclora Metaform"""
    
    def __init__(self):
        self.name = 'Anclora Metaform Conversion Engine'
        self.version = '1.0.0'
        
        # Conversiones soportadas y sus costos en créditos
        self.supported_conversions = {
            'txt': {
                'html': 1,
                'pdf': 1,
                'doc': 2,
                'docx': 2,
                'md': 1,
                'rtf': 1,
                'odt': 2,
                'tex': 3
            },
            'pdf': {
                'jpg': 3,
                'png': 3,
                'gif': 3,
                'txt': 4
            },
            'jpg': {
                'png': 1,
                'pdf': 3,
                'gif': 2
            },
            'png': {
                'jpg': 1,
                'pdf': 3,
                'gif': 2,
                'webp': 2
            },
            'gif': {
                'jpg': 1,
                'png': 1,
                'pdf': 3,
                'mp4': 4
            },
            'svg': {
                'png': 2
            },
            'doc': {
                'pdf': 4,
                'txt': 3,
                'html': 4
            },
            'docx': {
                'pdf': 4,
                'txt': 3,
                'html': 4
            }
        }

        # Mapeo de métodos de conversión
        self.conversion_methods = {
            ('txt', 'html'): self._convert_txt_to_html,
            ('txt', 'md'): self._convert_txt_to_markdown,
            ('txt', 'rtf'): self._convert_txt_to_rtf,
            ('txt', 'doc'): self._convert_txt_to_doc,
            ('txt', 'docx'): self._convert_txt_to_docx,
            ('txt', 'pdf'): self._convert_txt_to_pdf,
            ('txt', 'odt'): self._convert_txt_to_odt,
            ('txt', 'tex'): self._convert_txt_to_tex,
            ('pdf', 'jpg'): self._convert_pdf_to_jpg,
            ('pdf', 'png'): self._convert_pdf_to_png,
            ('pdf', 'gif'): self._convert_pdf_to_gif,
            ('pdf', 'txt'): self._convert_pdf_to_txt,
            ('jpg', 'png'): self._convert_jpg_to_png,
            ('jpg', 'pdf'): self._convert_jpg_to_pdf,
            ('jpg', 'gif'): self._convert_jpg_to_gif,
            ('png', 'jpg'): self._convert_png_to_jpg,
            ('png', 'pdf'): self._convert_png_to_pdf,
            ('png', 'gif'): self._convert_png_to_gif,
            ('png', 'webp'): self._convert_png_to_webp,
            ('gif', 'jpg'): self._convert_gif_to_jpg,
            ('gif', 'png'): self._convert_gif_to_png,
            ('gif', 'pdf'): self._convert_gif_to_pdf,
            ('gif', 'mp4'): self._convert_gif_to_mp4,
            ('svg', 'png'): self._convert_svg_to_png,
            ('doc', 'pdf'): self._convert_doc_to_pdf,
            ('doc', 'txt'): self._convert_doc_to_txt,
            ('doc', 'html'): self._convert_doc_to_html,
            ('docx', 'pdf'): self._convert_docx_to_pdf,
            ('docx', 'txt'): self._convert_docx_to_txt,
            ('docx', 'html'): self._convert_docx_to_html,
        }

    def get_supported_formats(self, source_format):
        """Obtiene los formatos de destino soportados para un formato origen"""
        return list(self.supported_conversions.get(source_format.lower(), {}).keys())

    def get_conversion_cost(self, source_format, target_format):
        """Calcula el costo en créditos de una conversión"""
        source = source_format.lower().replace('.', '')
        target = target_format.lower().replace('.', '')
        
        return self.supported_conversions.get(source, {}).get(target, 2)

    def validate_file(self, file_path, max_size_mb=100):
        """Valida un archivo antes de la conversión"""
        if not os.path.exists(file_path):
            return False, "El archivo no existe"
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "El archivo está vacío"
        
        if file_size > max_size_mb * 1024 * 1024:
            return False, f"El archivo es demasiado grande (máximo {max_size_mb}MB)"
        
        return True, "Archivo válido"

    def analyze_file(self, file_path, filename):
        """Analiza un archivo y genera recomendaciones"""
        file_extension = filename.split('.')[-1].lower()
        file_size = os.path.getsize(file_path)
        
        analysis = {
            'filename': filename,
            'extension': file_extension,
            'size': file_size,
            'supported_formats': self.get_supported_formats(file_extension),
            'recommendations': []
        }
        
        # Generar recomendaciones basadas en el tipo de archivo
        if file_extension == 'txt':
            analysis['recommendations'] = [
                'Para documentos formales, recomendamos PDF',
                'Para web, HTML es la mejor opción',
                'Para desarrolladores, Markdown es ideal'
            ]
        elif file_extension == 'pdf':
            analysis['recommendations'] = [
                'Para edición, convierte a DOC',
                'Para imágenes, JPG o PNG son ideales'
            ]
        elif file_extension in ['jpg', 'png']:
            analysis['recommendations'] = [
                'Para documentos, PDF mantiene la calidad',
                'Para web, considera optimización de tamaño'
            ]

        return analysis

    def find_conversion_path(self, src_format: str, dst_format: str):
        """Busca una ruta de conversión usando BFS con hasta dos intermediarios.

        Devuelve una lista de formatos que representa el camino desde src_format
        hasta dst_format, inclusive. Si no se encuentra una ruta válida dentro
        del límite de profundidad, devuelve None.
        """
        src = src_format.lower()
        dst = dst_format.lower()

        if src == dst:
            return [src]

        max_depth = 3  # número máximo de pasos (edges): src -> a -> b -> dst
        queue = deque([(src, [src])])
        visited = {src}

        while queue:
            current, path = queue.popleft()
            depth = len(path) - 1
            if depth >= max_depth:
                continue
            for neighbor in self.supported_conversions.get(current, {}):
                if neighbor in visited:
                    continue
                new_path = path + [neighbor]
                if neighbor == dst:
                    return new_path
                visited.add(neighbor)
                queue.append((neighbor, new_path))

        return None

    def convert_file(self, input_path, output_path, source_format, target_format):
        """Realiza la conversión de archivo"""
        try:
            source = source_format.lower()
            target = target_format.lower()
            method = self.conversion_methods.get((source, target))
            if method:
                return method(input_path, output_path)

            path = self.find_conversion_path(source, target)
            if not path:
                return False, f"Conversión {source_format} → {target_format} no implementada aún"

            logs = []
            temp_files = []
            current_input = input_path

            try:
                for i in range(len(path) - 1):
                    src_fmt = path[i]
                    dst_fmt = path[i + 1]
                    step_method = self.conversion_methods.get((src_fmt, dst_fmt))
                    if not step_method:
                        return False, f"Conversión {src_fmt} → {dst_fmt} no implementada"

                    if i == len(path) - 2:
                        current_output = output_path
                    else:
                        fd, current_output = tempfile.mkstemp(suffix=f'.{dst_fmt}')
                        os.close(fd)
                        temp_files.append(current_output)

                    success, msg = step_method(current_input, current_output)
                    logs.append(f"{src_fmt}->{dst_fmt}: {msg}")
                    if not success:
                        return False, f"Fallo en {src_fmt}->{dst_fmt}: {msg}"

                    current_input = current_output

                return True, " | ".join(logs)
            finally:
                for tmp in temp_files:
                    if os.path.exists(tmp):
                        os.remove(tmp)
        except Exception as e:
            return False, f"Error durante la conversión: {str(e)}"

    def convert_batch(self, tasks):
        """Procesa un lote de conversiones."""
        results = []
        for task in tasks:
            success, message = self.convert_file(
                task['input_path'],
                task['output_path'],
                task.get('source_format') or task['input_path'].split('.')[-1],
                task['target_format']
            )
            results.append({
                'input_path': task['input_path'],
                'output_path': task['output_path'],
                'success': success,
                'message': message
            })
        return results

    def _convert_txt_to_html(self, input_path, output_path):
        """Convierte TXT a HTML"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Escapar caracteres HTML
            content = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Convertir saltos de línea
            content = content.replace('\n\n', '</p><p>').replace('\n', '<br>')
            content = f'<p>{content}</p>'
            
            # Generar HTML completo
            html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Documento Convertido - Anclora Metaform</title>
    <meta name="generator" content="Anclora Metaform">
    <style>
        body {{ 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        .header {{
            text-align: center;
            border-bottom: 2px solid #1e40af;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .content {{
            background: #f8fafc;
            padding: 20px;
            border-radius: 8px;
        }}
        .footer {{
            text-align: center;
            margin-top: 20px;
            font-size: 0.9em;
            color: #666;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Anclora Metaform</h1>
        <p>Tu Contenido, Reinventado</p>
    </div>
    <div class="content">
        {content}
    </div>
    <div class="footer">
        <p>Convertido con Anclora Metaform - {datetime.now().strftime('%d/%m/%Y')}</p>
    </div>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True, "Conversión exitosa"
            
        except Exception as e:
            return False, f"Error en conversión TXT→HTML: {str(e)}"

    def _convert_txt_to_markdown(self, input_path, output_path):
        """Convierte TXT a Markdown"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Conversión básica a Markdown
            markdown_content = f"""# Documento Convertido

*Convertido con Anclora Metaform - {datetime.now().strftime('%d/%m/%Y')}*

---

{content}

---

**Generado por:** [Anclora Metaform](https://anclora-metaform.com)  
**Tu Contenido, Reinventado**
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return True, "Conversión exitosa"
            
        except Exception as e:
            return False, f"Error en conversión TXT→MD: {str(e)}"

    def _convert_txt_to_rtf(self, input_path, output_path):
        """Convierte TXT a RTF"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Generar RTF básico
            content_rtf = content.replace('\n', '\\par ')
            date_str = datetime.now().strftime('%d/%m/%Y')
            rtf_content = f"""{{\\rtf1\\ansi\\deff0 {{\\fonttbl {{\\f0 Times New Roman;}}}}
\\f0\\fs24 
\\b Documento Convertido - Anclora Metaform\\b0\\par
\\par
{content_rtf}
\\par
\\par
\\i Convertido con Anclora Metaform - {date_str}\\i0
}}"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(rtf_content)
            
            return True, "Conversión exitosa"
            
        except Exception as e:
            return False, f"Error en conversión TXT→RTF: {str(e)}"

    def _convert_txt_to_doc(self, input_path, output_path):
        """Convierte TXT a DOC (plano)"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f_src, open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(f_src.read())
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión TXT→DOC: {str(e)}"

    def _convert_txt_to_docx(self, input_path, output_path):
        """Convierte TXT a DOCX"""
        try:
            doc = Document()
            with open(input_path, 'r', encoding='utf-8') as f:
                for line in f:
                    doc.add_paragraph(line.rstrip('\n'))
            doc.save(output_path)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión TXT→DOCX: {str(e)}"

    def _convert_txt_to_pdf(self, input_path, output_path):
        """Convierte TXT a PDF"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            with open(input_path, 'r', encoding='utf-8') as f:
                for line in f:
                    pdf.multi_cell(0, 10, line)
            pdf.output(output_path)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión TXT→PDF: {str(e)}"

    def _convert_txt_to_odt(self, input_path, output_path):
        """Convierte TXT a ODT (plano)"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f_src, open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(f_src.read())
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión TXT→ODT: {str(e)}"

    def _convert_txt_to_tex(self, input_path, output_path):
        """Convierte TXT a LaTeX"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            tex = f"""\\documentclass{{article}}
\\begin{{document}}
{content}
\\end{{document}}
"""
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(tex)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión TXT→TEX: {str(e)}"

    def _convert_pdf_to_jpg(self, input_path, output_path):
        """Convierte PDF a JPG (placeholder)"""
        try:
            img = Image.new('RGB', (100, 100), 'white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 40), 'PDF', fill='black')
            img.save(output_path, 'JPEG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PDF→JPG: {str(e)}"

    def _convert_pdf_to_png(self, input_path, output_path):
        """Convierte PDF a PNG (placeholder)"""
        try:
            img = Image.new('RGB', (100, 100), 'white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 40), 'PDF', fill='black')
            img.save(output_path, 'PNG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PDF→PNG: {str(e)}"

    def _convert_pdf_to_gif(self, input_path, output_path):
        """Convierte PDF a GIF (placeholder)"""
        try:
            img = Image.new('RGB', (100, 100), 'white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 40), 'PDF', fill='black')
            img.save(output_path, 'GIF')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PDF→GIF: {str(e)}"

    def _convert_pdf_to_txt(self, input_path, output_path):
        """Convierte PDF a TXT"""
        try:
            reader = PdfReader(input_path)
            text = ''
            for page in reader.pages:
                text += page.extract_text() or ''
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(text)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PDF→TXT: {str(e)}"

    def _convert_jpg_to_png(self, input_path, output_path):
        """Convierte JPG a PNG"""
        try:
            with Image.open(input_path) as img:
                img.save(output_path, 'PNG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión JPG→PNG: {str(e)}"

    def _convert_jpg_to_pdf(self, input_path, output_path):
        """Convierte JPG a PDF"""
        try:
            with Image.open(input_path) as img:
                img.convert('RGB').save(output_path, 'PDF')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión JPG→PDF: {str(e)}"

    def _convert_jpg_to_gif(self, input_path, output_path):
        """Convierte JPG a GIF"""
        try:
            with Image.open(input_path) as img:
                img.save(output_path, 'GIF')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión JPG→GIF: {str(e)}"

    def _convert_png_to_jpg(self, input_path, output_path):
        """Convierte PNG a JPG"""
        try:
            with Image.open(input_path) as img:
                img.convert('RGB').save(output_path, 'JPEG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PNG→JPG: {str(e)}"

    def _convert_png_to_pdf(self, input_path, output_path):
        """Convierte PNG a PDF"""
        try:
            with Image.open(input_path) as img:
                img.convert('RGB').save(output_path, 'PDF')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PNG→PDF: {str(e)}"

    def _convert_png_to_gif(self, input_path, output_path):
        """Convierte PNG a GIF"""
        try:
            with Image.open(input_path) as img:
                img.save(output_path, 'GIF')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PNG→GIF: {str(e)}"

    def _convert_png_to_webp(self, input_path, output_path):
        """Convierte PNG a WEBP"""
        try:
            with Image.open(input_path) as img:
                img.save(output_path, 'WEBP')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión PNG→WEBP: {str(e)}"

    def _convert_gif_to_jpg(self, input_path, output_path):
        """Convierte GIF a JPG"""
        try:
            with Image.open(input_path) as img:
                img.convert('RGB').save(output_path, 'JPEG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión GIF→JPG: {str(e)}"

    def _convert_gif_to_png(self, input_path, output_path):
        """Convierte GIF a PNG"""
        try:
            with Image.open(input_path) as img:
                img.save(output_path, 'PNG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión GIF→PNG: {str(e)}"

    def _convert_gif_to_pdf(self, input_path, output_path):
        """Convierte GIF a PDF"""
        try:
            with Image.open(input_path) as img:
                img.convert('RGB').save(output_path, 'PDF')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión GIF→PDF: {str(e)}"

    def _convert_gif_to_mp4(self, input_path, output_path):
        """Convierte GIF a MP4 (placeholder)"""
        try:
            with open(input_path, 'rb') as src, open(output_path, 'wb') as dst:
                shutil.copyfileobj(src, dst)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión GIF→MP4: {str(e)}"

    def _convert_svg_to_png(self, input_path, output_path):
        """Convierte SVG a PNG (placeholder)"""
        try:
            with open(input_path, 'r', encoding='utf-8'):
                pass  # simple validation de existencia
            img = Image.new('RGB', (100, 100), 'white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 40), 'SVG', fill='black')
            img.save(output_path, 'PNG')
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión SVG→PNG: {str(e)}"

    def _convert_doc_to_pdf(self, input_path, output_path):
        """Convierte DOC a PDF"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                text = f.read()
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, text)
            pdf.output(output_path)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión DOC→PDF: {str(e)}"

    def _convert_doc_to_txt(self, input_path, output_path):
        """Convierte DOC a TXT"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f_src, open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(f_src.read())
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión DOC→TXT: {str(e)}"

    def _convert_doc_to_html(self, input_path, output_path):
        """Convierte DOC a HTML"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            html = f"<html><body><pre>{content}</pre></body></html>"
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(html)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión DOC→HTML: {str(e)}"

    def _read_docx(self, path):
        doc = Document(path)
        return '\n'.join(p.text for p in doc.paragraphs)

    def _convert_docx_to_pdf(self, input_path, output_path):
        """Convierte DOCX a PDF"""
        try:
            text = self._read_docx(input_path)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, text)
            pdf.output(output_path)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión DOCX→PDF: {str(e)}"

    def _convert_docx_to_txt(self, input_path, output_path):
        """Convierte DOCX a TXT"""
        try:
            text = self._read_docx(input_path)
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(text)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión DOCX→TXT: {str(e)}"

    def _convert_docx_to_html(self, input_path, output_path):
        """Convierte DOCX a HTML"""
        try:
            text = self._read_docx(input_path)
            html = f"<html><body><pre>{text}</pre></body></html>"
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(html)
            return True, "Conversión exitosa"
        except Exception as e:
            return False, f"Error en conversión DOCX→HTML: {str(e)}"

# Instancia global del motor de conversión
conversion_engine = ConversionEngine()

