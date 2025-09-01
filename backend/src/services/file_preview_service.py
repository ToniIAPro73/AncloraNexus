# ================================
# ANCLORA NEXUS - FILE PREVIEW SERVICE
# Sistema de vista previa de archivos
# ================================

import os
import base64
import tempfile
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from PIL import Image, ImageDraw, ImageFont
import io

@dataclass
class PreviewData:
    """Datos de vista previa de archivo"""
    preview_type: str  # 'text', 'image', 'thumbnail', 'metadata'
    content: str  # Base64 para imágenes, texto para documentos
    metadata: Dict[str, Any]
    page_count: Optional[int] = None
    dimensions: Optional[Tuple[int, int]] = None
    file_info: Optional[Dict[str, Any]] = None
    preview_quality: str = 'medium'  # 'low', 'medium', 'high'

class FilePreviewService:
    """Servicio de vista previa de archivos"""
    
    def __init__(self):
        self.max_text_preview = 2000  # Caracteres máximos para preview de texto
        self.thumbnail_size = (300, 400)  # Tamaño de thumbnail
        self.supported_formats = {
            'text': ['txt', 'md', 'html', 'rtf', 'csv'],
            'image': ['jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp'],
            'document': ['pdf', 'doc', 'docx', 'odt'],
            'code': ['py', 'js', 'ts', 'css', 'json', 'xml', 'yaml']
        }

    async def generate_preview(self, file_path: str, filename: str, 
                             preview_quality: str = 'medium') -> PreviewData:
        """Genera vista previa del archivo"""
        
        file_extension = filename.split('.')[-1].lower()
        file_stats = os.stat(file_path)
        
        # Información básica del archivo
        file_info = {
            'name': filename,
            'size': file_stats.st_size,
            'size_mb': file_stats.st_size / (1024 * 1024),
            'extension': file_extension,
            'created': file_stats.st_ctime,
            'modified': file_stats.st_mtime
        }
        
        # Generar preview según el tipo de archivo
        if file_extension in self.supported_formats['text']:
            return await self._generate_text_preview(file_path, filename, file_info, preview_quality)
        
        elif file_extension in self.supported_formats['image']:
            return await self._generate_image_preview(file_path, filename, file_info, preview_quality)
        
        elif file_extension in self.supported_formats['document']:
            return await self._generate_document_preview(file_path, filename, file_info, preview_quality)
        
        elif file_extension in self.supported_formats['code']:
            return await self._generate_code_preview(file_path, filename, file_info, preview_quality)
        
        else:
            return await self._generate_generic_preview(file_path, filename, file_info)

    async def _generate_text_preview(self, file_path: str, filename: str, 
                                   file_info: Dict, quality: str) -> PreviewData:
        """Genera preview para archivos de texto"""
        
        try:
            # Leer contenido del archivo
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(self.max_text_preview)
            
            # Contar líneas y palabras
            lines = content.count('\n') + 1
            words = len(content.split())
            
            metadata = {
                'lines': lines,
                'words': words,
                'characters': len(content),
                'encoding': 'utf-8',
                'preview_truncated': len(content) >= self.max_text_preview
            }
            
            return PreviewData(
                preview_type='text',
                content=content,
                metadata=metadata,
                file_info=file_info,
                preview_quality=quality
            )
            
        except Exception as e:
            return PreviewData(
                preview_type='text',
                content=f"Error leyendo archivo: {str(e)}",
                metadata={'error': str(e)},
                file_info=file_info,
                preview_quality='low'
            )

    async def _generate_image_preview(self, file_path: str, filename: str, 
                                    file_info: Dict, quality: str) -> PreviewData:
        """Genera preview para imágenes"""
        
        try:
            with Image.open(file_path) as img:
                # Información de la imagen
                metadata = {
                    'width': img.width,
                    'height': img.height,
                    'mode': img.mode,
                    'format': img.format,
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                }
                
                # Generar thumbnail
                thumbnail_size = self.thumbnail_size
                if quality == 'high':
                    thumbnail_size = (600, 800)
                elif quality == 'low':
                    thumbnail_size = (150, 200)
                
                img.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
                
                # Convertir a base64
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                
                return PreviewData(
                    preview_type='image',
                    content=img_base64,
                    metadata=metadata,
                    dimensions=(img.width, img.height),
                    file_info=file_info,
                    preview_quality=quality
                )
                
        except Exception as e:
            return await self._generate_error_thumbnail(filename, str(e), file_info)

    async def _generate_document_preview(self, file_path: str, filename: str, 
                                       file_info: Dict, quality: str) -> PreviewData:
        """Genera preview para documentos"""
        
        file_extension = filename.split('.')[-1].lower()
        
        try:
            if file_extension == 'pdf':
                return await self._generate_pdf_preview(file_path, filename, file_info, quality)
            elif file_extension in ['doc', 'docx']:
                return await self._generate_word_preview(file_path, filename, file_info, quality)
            else:
                return await self._generate_generic_preview(file_path, filename, file_info)
                
        except Exception as e:
            return await self._generate_error_thumbnail(filename, str(e), file_info)

    async def _generate_pdf_preview(self, file_path: str, filename: str, 
                                  file_info: Dict, quality: str) -> PreviewData:
        """Genera preview específico para PDF"""
        
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            
            # Extraer texto de la primera página
            first_page_text = ""
            if reader.pages:
                first_page_text = reader.pages[0].extract_text()[:self.max_text_preview]
            
            # Metadata del PDF
            metadata = {
                'page_count': page_count,
                'title': reader.metadata.get('/Title', '') if reader.metadata else '',
                'author': reader.metadata.get('/Author', '') if reader.metadata else '',
                'creator': reader.metadata.get('/Creator', '') if reader.metadata else '',
                'has_text': len(first_page_text.strip()) > 0,
                'is_encrypted': reader.is_encrypted
            }
            
            # Generar thumbnail del PDF (simulado)
            thumbnail = await self._generate_pdf_thumbnail(first_page_text, filename)
            
            return PreviewData(
                preview_type='document',
                content=thumbnail,  # Base64 del thumbnail
                metadata=metadata,
                page_count=page_count,
                file_info=file_info,
                preview_quality=quality
            )
            
        except Exception as e:
            return await self._generate_error_thumbnail(filename, str(e), file_info)

    async def _generate_word_preview(self, file_path: str, filename: str, 
                                   file_info: Dict, quality: str) -> PreviewData:
        """Genera preview para documentos Word"""
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extraer texto de los primeros párrafos
            text_content = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if len(text_content) < self.max_text_preview:
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1
                else:
                    break
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'preview_paragraphs': paragraph_count,
                'has_images': len(doc.inline_shapes) > 0,
                'has_tables': len(doc.tables) > 0,
                'preview_truncated': len(text_content) >= self.max_text_preview
            }
            
            return PreviewData(
                preview_type='text',
                content=text_content,
                metadata=metadata,
                file_info=file_info,
                preview_quality=quality
            )
            
        except Exception as e:
            return await self._generate_error_thumbnail(filename, str(e), file_info)

    async def _generate_code_preview(self, file_path: str, filename: str, 
                                   file_info: Dict, quality: str) -> PreviewData:
        """Genera preview para archivos de código"""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(self.max_text_preview)
            
            # Análisis básico del código
            lines = content.split('\n')
            metadata = {
                'lines': len(lines),
                'language': filename.split('.')[-1].lower(),
                'estimated_functions': content.count('def ') + content.count('function '),
                'estimated_classes': content.count('class '),
                'has_comments': '//' in content or '#' in content or '/*' in content
            }
            
            return PreviewData(
                preview_type='code',
                content=content,
                metadata=metadata,
                file_info=file_info,
                preview_quality=quality
            )
            
        except Exception as e:
            return await self._generate_error_thumbnail(filename, str(e), file_info)

    async def _generate_pdf_thumbnail(self, text_content: str, filename: str) -> str:
        """Genera thumbnail visual para PDF"""
        
        try:
            # Crear imagen thumbnail simulando página PDF
            img = Image.new('RGB', (300, 400), color='white')
            draw = ImageDraw.Draw(img)
            
            # Dibujar borde de página
            draw.rectangle([10, 10, 290, 390], outline='#cccccc', width=2)
            
            # Agregar texto simulado
            try:
                # Intentar usar fuente del sistema
                font = ImageFont.truetype("arial.ttf", 12)
            except:
                font = ImageFont.load_default()
            
            # Dibujar líneas de texto
            y_position = 25
            for line in text_content[:200].split('\n')[:15]:  # Primeras 15 líneas
                if y_position > 360:
                    break
                draw.text((20, y_position), line[:40], fill='black', font=font)
                y_position += 20
            
            # Agregar indicador de PDF
            draw.text((20, 370), f"📄 {filename}", fill='#666666', font=font)
            
            # Convertir a base64
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            return base64.b64encode(buffer.getvalue()).decode('utf-8')
            
        except Exception as e:
            print(f"Error generating PDF thumbnail: {e}")
            return ""

    async def _generate_error_thumbnail(self, filename: str, error: str, file_info: Dict) -> PreviewData:
        """Genera thumbnail de error"""
        
        try:
            img = Image.new('RGB', (300, 400), color='#f8f8f8')
            draw = ImageDraw.Draw(img)
            
            # Dibujar borde de error
            draw.rectangle([10, 10, 290, 390], outline='#ff6b6b', width=2)
            
            # Texto de error
            try:
                font = ImageFont.truetype("arial.ttf", 14)
                small_font = ImageFont.truetype("arial.ttf", 10)
            except:
                font = ImageFont.load_default()
                small_font = ImageFont.load_default()
            
            draw.text((20, 50), "❌ Error de Preview", fill='#ff6b6b', font=font)
            draw.text((20, 80), filename[:30], fill='black', font=small_font)
            draw.text((20, 100), f"Tamaño: {file_info['size_mb']:.1f} MB", fill='#666', font=small_font)
            
            # Convertir a base64
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            
            return PreviewData(
                preview_type='thumbnail',
                content=img_base64,
                metadata={'error': error, 'preview_failed': True},
                file_info=file_info,
                preview_quality='low'
            )
            
        except Exception:
            return PreviewData(
                preview_type='metadata',
                content="Preview no disponible",
                metadata={'error': error},
                file_info=file_info,
                preview_quality='low'
            )

    async def _generate_generic_preview(self, file_path: str, filename: str, file_info: Dict) -> PreviewData:
        """Genera preview genérico para formatos no soportados"""
        
        file_extension = filename.split('.')[-1].lower()
        
        # Crear thumbnail genérico
        try:
            img = Image.new('RGB', (300, 400), color='#f0f0f0')
            draw = ImageDraw.Draw(img)
            
            # Dibujar icono de archivo
            draw.rectangle([100, 100, 200, 250], outline='#4a90e2', width=3, fill='#e3f2fd')
            draw.polygon([(200, 100), (200, 130), (170, 100)], fill='#4a90e2')
            
            # Texto del archivo
            try:
                font = ImageFont.truetype("arial.ttf", 16)
                small_font = ImageFont.truetype("arial.ttf", 12)
            except:
                font = ImageFont.load_default()
                small_font = ImageFont.load_default()
            
            draw.text((120, 180), file_extension.upper(), fill='#4a90e2', font=font)
            draw.text((50, 280), filename[:25], fill='black', font=small_font)
            draw.text((50, 300), f"{file_info['size_mb']:.1f} MB", fill='#666', font=small_font)
            
            # Convertir a base64
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            
            metadata = {
                'format': file_extension,
                'preview_type': 'generic_thumbnail',
                'supported_preview': False
            }
            
            return PreviewData(
                preview_type='thumbnail',
                content=img_base64,
                metadata=metadata,
                file_info=file_info,
                preview_quality='medium'
            )
            
        except Exception as e:
            return PreviewData(
                preview_type='metadata',
                content=f"Archivo {filename}",
                metadata={'error': str(e)},
                file_info=file_info,
                preview_quality='low'
            )

    def get_preview_capabilities(self, filename: str) -> Dict[str, Any]:
        """Obtiene capacidades de preview para un archivo"""
        
        file_extension = filename.split('.')[-1].lower()
        
        capabilities = {
            'can_preview': False,
            'preview_type': 'none',
            'features': [],
            'limitations': []
        }
        
        if file_extension in self.supported_formats['text']:
            capabilities.update({
                'can_preview': True,
                'preview_type': 'text',
                'features': ['Contenido completo', 'Conteo de palabras', 'Análisis de codificación'],
                'limitations': [f'Limitado a {self.max_text_preview} caracteres']
            })
        
        elif file_extension in self.supported_formats['image']:
            capabilities.update({
                'can_preview': True,
                'preview_type': 'image',
                'features': ['Thumbnail de alta calidad', 'Información EXIF', 'Dimensiones'],
                'limitations': ['Redimensionado para web']
            })
        
        elif file_extension in self.supported_formats['document']:
            capabilities.update({
                'can_preview': True,
                'preview_type': 'document',
                'features': ['Primera página', 'Metadatos', 'Conteo de páginas'],
                'limitations': ['Solo primera página', 'Sin formato visual completo']
            })
        
        elif file_extension in self.supported_formats['code']:
            capabilities.update({
                'can_preview': True,
                'preview_type': 'code',
                'features': ['Sintaxis resaltada', 'Análisis de estructura', 'Conteo de líneas'],
                'limitations': ['Sin ejecución de código']
            })
        
        else:
            capabilities.update({
                'can_preview': True,
                'preview_type': 'generic',
                'features': ['Información básica', 'Thumbnail genérico'],
                'limitations': ['Preview limitado para este formato']
            })
        
        return capabilities

# Instancia global del servicio de preview
file_preview_service = FilePreviewService()
