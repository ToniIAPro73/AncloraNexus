# src/services/ai_content_analyzer.py - Análisis inteligente de contenido de documentos
import os
import re
import json
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import hashlib

# Análisis de contenido
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    print("Warning: python-magic no disponible, usando detección básica de tipos")

from PIL import Image
import PyPDF2
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET

@dataclass
class ContentAnalysis:
    """Resultado del análisis de contenido"""
    file_type: str
    content_type: str
    complexity_score: float
    quality_score: float
    size_mb: float
    page_count: int
    word_count: int
    image_count: int
    table_count: int
    has_forms: bool
    has_hyperlinks: bool
    has_embedded_media: bool
    text_to_image_ratio: float
    recommended_formats: List[Dict[str, Any]]
    optimization_suggestions: List[str]
    quality_issues: List[str]
    metadata: Dict[str, Any]

class AIContentAnalyzer:
    """Analizador inteligente de contenido de documentos"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._analyze_pdf,
            'docx': self._analyze_docx,
            'doc': self._analyze_docx,  # Usar el mismo método para DOC
            'txt': self._analyze_text,
            'html': self._analyze_text,  # Usar análisis de texto para HTML
            'xml': self._analyze_text,   # Usar análisis de texto para XML
            'csv': self._analyze_text,   # Usar análisis de texto para CSV
            'xlsx': self._analyze_text,  # Simplificado por ahora
            'png': self._analyze_image,
            'jpg': self._analyze_image,
            'jpeg': self._analyze_image,
            'gif': self._analyze_image,
            'bmp': self._analyze_image,
            'webp': self._analyze_image,
            'svg': self._analyze_text    # Tratar SVG como texto
        }
        
        # Reglas de recomendación basadas en contenido
        self.format_rules = {
            'high_quality_images': ['pdf', 'png', 'tiff'],
            'web_display': ['html', 'webp', 'svg'],
            'print_ready': ['pdf', 'eps', 'tiff'],
            'data_analysis': ['csv', 'xlsx', 'json'],
            'archival': ['pdf/a', 'tiff', 'png'],
            'mobile_friendly': ['webp', 'html', 'epub'],
            'interactive': ['html', 'pdf', 'epub'],
            'lightweight': ['webp', 'svg', 'txt', 'md']
        }

    def analyze_file(self, file_path: str) -> ContentAnalysis:
        """Analizar archivo y generar recomendaciones inteligentes"""
        try:
            file_path = Path(file_path)
            
            # Información básica del archivo
            file_size = file_path.stat().st_size / (1024 * 1024)  # MB
            file_ext = file_path.suffix.lower().lstrip('.')
            
            # Detectar tipo MIME
            if MAGIC_AVAILABLE:
                mime_type = magic.from_file(str(file_path), mime=True)
            else:
                mime_type = mimetypes.guess_type(str(file_path))[0] or 'application/octet-stream'
            
            # Análisis específico por formato
            if file_ext in self.supported_formats:
                analysis_data = self.supported_formats[file_ext](file_path)
            else:
                analysis_data = self._analyze_generic(file_path)
            
            # Calcular scores de complejidad y calidad
            complexity_score = self._calculate_complexity_score(analysis_data)
            quality_score = self._calculate_quality_score(analysis_data, file_path)
            
            # Generar recomendaciones
            recommendations = self._generate_format_recommendations(
                analysis_data, complexity_score, quality_score, file_ext
            )
            
            # Sugerencias de optimización
            optimizations = self._generate_optimization_suggestions(
                analysis_data, complexity_score, quality_score
            )
            
            # Detectar problemas de calidad
            quality_issues = self._detect_quality_issues(analysis_data, file_path)
            
            return ContentAnalysis(
                file_type=file_ext,
                content_type=analysis_data.get('content_type', 'unknown'),
                complexity_score=complexity_score,
                quality_score=quality_score,
                size_mb=file_size,
                page_count=analysis_data.get('page_count', 1),
                word_count=analysis_data.get('word_count', 0),
                image_count=analysis_data.get('image_count', 0),
                table_count=analysis_data.get('table_count', 0),
                has_forms=analysis_data.get('has_forms', False),
                has_hyperlinks=analysis_data.get('has_hyperlinks', False),
                has_embedded_media=analysis_data.get('has_embedded_media', False),
                text_to_image_ratio=analysis_data.get('text_to_image_ratio', 1.0),
                recommended_formats=recommendations,
                optimization_suggestions=optimizations,
                quality_issues=quality_issues,
                metadata=analysis_data.get('metadata', {})
            )
            
        except Exception as e:
            # Análisis básico en caso de error
            return self._create_fallback_analysis(file_path, str(e))

    def _analyze_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Análisis específico para archivos PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                page_count = len(pdf_reader.pages)
                text_content = ""
                image_count = 0
                has_forms = False
                has_hyperlinks = False
                
                # Extraer texto y contar elementos
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
                    
                    # Detectar imágenes (aproximado)
                    if '/XObject' in page.get('/Resources', {}):
                        xobjects = page['/Resources']['/XObject'].get_object()
                        for obj in xobjects:
                            if xobjects[obj]['/Subtype'] == '/Image':
                                image_count += 1
                    
                    # Detectar formularios
                    if '/Annots' in page:
                        has_forms = True
                    
                    # Detectar enlaces
                    if '/Link' in str(page):
                        has_hyperlinks = True
                
                word_count = len(text_content.split())
                
                # Detectar tipo de contenido
                content_type = self._classify_content_type(text_content, image_count, page_count)
                
                return {
                    'content_type': content_type,
                    'page_count': page_count,
                    'word_count': word_count,
                    'image_count': image_count,
                    'table_count': text_content.count('\t') // 10,  # Aproximado
                    'has_forms': has_forms,
                    'has_hyperlinks': has_hyperlinks,
                    'has_embedded_media': image_count > 0,
                    'text_to_image_ratio': word_count / max(image_count, 1),
                    'metadata': {
                        'creator': pdf_reader.metadata.get('/Creator', 'Unknown') if pdf_reader.metadata else 'Unknown',
                        'producer': pdf_reader.metadata.get('/Producer', 'Unknown') if pdf_reader.metadata else 'Unknown',
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')) if pdf_reader.metadata else ''
                    }
                }
                
        except Exception as e:
            return {'content_type': 'document', 'error': str(e)}

    def _analyze_docx(self, file_path: Path) -> Dict[str, Any]:
        """Análisis específico para archivos DOCX"""
        try:
            doc = Document(file_path)
            
            # Contar elementos
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Extraer texto
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + " "
            
            word_count = len(text_content.split())
            
            # Contar imágenes y elementos multimedia
            image_count = 0
            has_hyperlinks = False
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            # Detectar enlaces
            if "hyperlink" in text_content.lower():
                has_hyperlinks = True
            
            content_type = self._classify_content_type(text_content, image_count, 1)
            
            return {
                'content_type': content_type,
                'page_count': max(1, paragraph_count // 20),  # Estimado
                'word_count': word_count,
                'image_count': image_count,
                'table_count': table_count,
                'has_forms': False,
                'has_hyperlinks': has_hyperlinks,
                'has_embedded_media': image_count > 0,
                'text_to_image_ratio': word_count / max(image_count, 1),
                'metadata': {
                    'paragraphs': paragraph_count,
                    'tables': table_count
                }
            }
            
        except Exception as e:
            return {'content_type': 'document', 'error': str(e)}

    def _analyze_image(self, file_path: Path) -> Dict[str, Any]:
        """Análisis específico para imágenes"""
        try:
            with Image.open(file_path) as img:
                width, height = img.size
                mode = img.mode
                format_name = img.format
                
                # Calcular aspectos de calidad
                resolution = width * height
                aspect_ratio = width / height
                
                # Detectar tipo de contenido de imagen
                if aspect_ratio > 2 or aspect_ratio < 0.5:
                    content_type = 'banner'
                elif resolution > 2000000:  # > 2MP
                    content_type = 'high_resolution'
                elif width < 200 or height < 200:
                    content_type = 'icon'
                else:
                    content_type = 'standard_image'
                
                return {
                    'content_type': content_type,
                    'page_count': 1,
                    'word_count': 0,
                    'image_count': 1,
                    'table_count': 0,
                    'has_forms': False,
                    'has_hyperlinks': False,
                    'has_embedded_media': False,
                    'text_to_image_ratio': 0,
                    'metadata': {
                        'width': width,
                        'height': height,
                        'mode': mode,
                        'format': format_name,
                        'resolution': resolution,
                        'aspect_ratio': round(aspect_ratio, 2)
                    }
                }
                
        except Exception as e:
            return {'content_type': 'image', 'error': str(e)}

    def _analyze_text(self, file_path: Path) -> Dict[str, Any]:
        """Análisis específico para archivos de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            word_count = len(content.split())
            line_count = len(content.split('\n'))
            
            # Detectar tipo de contenido textual
            if content.count('#') > 5:  # Markdown
                content_type = 'markdown'
            elif content.count('<') > 10:  # HTML
                content_type = 'html'
            elif content.count('{') > 5:  # JSON/Code
                content_type = 'structured_data'
            else:
                content_type = 'plain_text'
            
            return {
                'content_type': content_type,
                'page_count': max(1, line_count // 50),
                'word_count': word_count,
                'image_count': 0,
                'table_count': content.count('|') // 10,  # Aproximado para tablas markdown
                'has_forms': False,
                'has_hyperlinks': 'http' in content.lower(),
                'has_embedded_media': False,
                'text_to_image_ratio': float('inf'),
                'metadata': {
                    'lines': line_count,
                    'characters': len(content)
                }
            }
            
        except Exception as e:
            return {'content_type': 'text', 'error': str(e)}

    def _analyze_generic(self, file_path: Path) -> Dict[str, Any]:
        """Análisis genérico para formatos no soportados específicamente"""
        file_size = file_path.stat().st_size
        
        return {
            'content_type': 'unknown',
            'page_count': 1,
            'word_count': 0,
            'image_count': 0,
            'table_count': 0,
            'has_forms': False,
            'has_hyperlinks': False,
            'has_embedded_media': False,
            'text_to_image_ratio': 1.0,
            'metadata': {
                'file_size': file_size,
                'extension': file_path.suffix
            }
        }

    def _classify_content_type(self, text_content: str, image_count: int, page_count: int) -> str:
        """Clasificar el tipo de contenido basado en características"""
        text_lower = text_content.lower()

        # Detectar documentos técnicos
        if any(word in text_lower for word in ['api', 'function', 'class', 'method', 'algorithm']):
            return 'technical_document'

        # Detectar reportes/informes
        if any(word in text_lower for word in ['report', 'analysis', 'summary', 'conclusion', 'findings']):
            return 'report'

        # Detectar presentaciones
        if page_count > 5 and len(text_content.split()) / page_count < 100:
            return 'presentation'

        # Detectar documentos con muchas imágenes
        if image_count > page_count * 2:
            return 'visual_document'

        # Detectar formularios
        if any(word in text_lower for word in ['form', 'application', 'submit', 'signature']):
            return 'form'

        # Detectar manuales
        if any(word in text_lower for word in ['manual', 'guide', 'instructions', 'tutorial']):
            return 'manual'

        return 'general_document'

    def _calculate_complexity_score(self, analysis_data: Dict[str, Any]) -> float:
        """Calcular score de complejidad del documento (0-100)"""
        score = 0

        # Complejidad por número de páginas
        page_count = analysis_data.get('page_count', 1)
        score += min(page_count * 2, 20)

        # Complejidad por elementos multimedia
        image_count = analysis_data.get('image_count', 0)
        score += min(image_count * 5, 25)

        # Complejidad por tablas
        table_count = analysis_data.get('table_count', 0)
        score += min(table_count * 3, 15)

        # Complejidad por formularios
        if analysis_data.get('has_forms', False):
            score += 15

        # Complejidad por enlaces
        if analysis_data.get('has_hyperlinks', False):
            score += 10

        # Complejidad por contenido embebido
        if analysis_data.get('has_embedded_media', False):
            score += 15

        return min(score, 100)

    def _calculate_quality_score(self, analysis_data: Dict[str, Any], file_path: Path) -> float:
        """Calcular score de calidad del documento (0-100)"""
        score = 100

        # Penalizar por errores
        if 'error' in analysis_data:
            score -= 30

        # Evaluar resolución de imágenes (si aplica)
        metadata = analysis_data.get('metadata', {})
        if 'resolution' in metadata:
            resolution = metadata['resolution']
            if resolution < 100000:  # Muy baja resolución
                score -= 20
            elif resolution > 10000000:  # Muy alta resolución (puede ser excesiva)
                score -= 10

        # Evaluar tamaño de archivo vs contenido
        file_size = file_path.stat().st_size
        word_count = analysis_data.get('word_count', 0)

        if word_count > 0:
            size_per_word = file_size / word_count
            if size_per_word > 1000:  # Archivo muy pesado para su contenido
                score -= 15

        # Evaluar estructura del documento
        if analysis_data.get('content_type') == 'unknown':
            score -= 10

        return max(score, 0)

    def _generate_format_recommendations(self, analysis_data: Dict[str, Any],
                                       complexity_score: float, quality_score: float,
                                       current_format: str) -> List[Dict[str, Any]]:
        """Generar recomendaciones de formato basadas en el análisis"""
        recommendations = []
        content_type = analysis_data.get('content_type', 'unknown')

        # Recomendaciones basadas en tipo de contenido
        if content_type == 'presentation':
            recommendations.extend([
                {'format': 'pdf', 'reason': 'Preserva el diseño y es universalmente compatible', 'priority': 'high'},
                {'format': 'html', 'reason': 'Permite interactividad y es fácil de compartir', 'priority': 'medium'},
                {'format': 'png', 'reason': 'Ideal para slides individuales de alta calidad', 'priority': 'low'}
            ])

        elif content_type == 'visual_document':
            recommendations.extend([
                {'format': 'pdf', 'reason': 'Mantiene la calidad de imágenes y el layout', 'priority': 'high'},
                {'format': 'html', 'reason': 'Optimiza imágenes para web', 'priority': 'medium'},
                {'format': 'epub', 'reason': 'Ideal para lectura en dispositivos móviles', 'priority': 'low'}
            ])

        elif content_type == 'technical_document':
            recommendations.extend([
                {'format': 'pdf', 'reason': 'Preserva formato técnico y es archivable', 'priority': 'high'},
                {'format': 'html', 'reason': 'Permite búsqueda y navegación fácil', 'priority': 'medium'},
                {'format': 'md', 'reason': 'Formato estándar para documentación técnica', 'priority': 'medium'}
            ])

        elif content_type == 'form':
            recommendations.extend([
                {'format': 'pdf', 'reason': 'Mantiene campos de formulario interactivos', 'priority': 'high'},
                {'format': 'html', 'reason': 'Permite validación y envío online', 'priority': 'high'},
                {'format': 'docx', 'reason': 'Fácil de editar y completar', 'priority': 'medium'}
            ])

        # Recomendaciones basadas en complejidad
        if complexity_score > 70:
            recommendations.append({
                'format': 'pdf',
                'reason': 'Documento complejo - PDF preserva mejor la estructura',
                'priority': 'high'
            })
        elif complexity_score < 30:
            recommendations.extend([
                {'format': 'txt', 'reason': 'Documento simple - formato ligero', 'priority': 'medium'},
                {'format': 'html', 'reason': 'Fácil visualización y compartir', 'priority': 'medium'}
            ])

        # Recomendaciones basadas en calidad
        if quality_score < 50:
            recommendations.append({
                'format': 'txt',
                'reason': 'Calidad baja detectada - formato simple para preservar contenido',
                'priority': 'medium'
            })

        # Filtrar formato actual y duplicados
        recommendations = [r for r in recommendations if r['format'] != current_format]
        seen_formats = set()
        unique_recommendations = []
        for rec in recommendations:
            if rec['format'] not in seen_formats:
                seen_formats.add(rec['format'])
                unique_recommendations.append(rec)

        # Ordenar por prioridad
        priority_order = {'high': 3, 'medium': 2, 'low': 1}
        unique_recommendations.sort(key=lambda x: priority_order.get(x['priority'], 0), reverse=True)

        return unique_recommendations[:5]  # Máximo 5 recomendaciones

    def _generate_optimization_suggestions(self, analysis_data: Dict[str, Any],
                                         complexity_score: float, quality_score: float) -> List[str]:
        """Generar sugerencias de optimización"""
        suggestions = []

        # Sugerencias basadas en tamaño
        if analysis_data.get('image_count', 0) > 10:
            suggestions.append("🖼️ Considera comprimir las imágenes para reducir el tamaño del archivo")

        # Sugerencias basadas en estructura
        if analysis_data.get('table_count', 0) > 5:
            suggestions.append("📊 Para documentos con muchas tablas, considera usar formato Excel o CSV")

        # Sugerencias basadas en contenido
        if analysis_data.get('has_hyperlinks', False):
            suggestions.append("🔗 Los enlaces se preservan mejor en formatos HTML o PDF")

        if analysis_data.get('has_forms', False):
            suggestions.append("📝 Los formularios interactivos funcionan mejor en PDF o HTML")

        # Sugerencias basadas en calidad
        if quality_score < 70:
            suggestions.append("⚠️ Se detectaron posibles problemas de calidad - revisa el archivo original")

        # Sugerencias basadas en complejidad
        if complexity_score > 80:
            suggestions.append("🔧 Documento muy complejo - considera dividirlo en secciones más pequeñas")

        # Sugerencias generales
        metadata = analysis_data.get('metadata', {})
        if 'resolution' in metadata and metadata['resolution'] > 5000000:
            suggestions.append("📐 Imágenes de muy alta resolución - considera reducir para uso web")

        return suggestions

    def _detect_quality_issues(self, analysis_data: Dict[str, Any], file_path: Path) -> List[str]:
        """Detectar problemas de calidad en el documento"""
        issues = []

        # Problemas de archivo
        if 'error' in analysis_data:
            issues.append(f"❌ Error al procesar el archivo: {analysis_data['error']}")

        # Problemas de tamaño
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > 50:
            issues.append("📏 Archivo muy grande - puede causar problemas de rendimiento")

        # Problemas de contenido
        word_count = analysis_data.get('word_count', 0)
        if word_count == 0 and analysis_data.get('image_count', 0) == 0:
            issues.append("📄 Documento parece estar vacío o no se pudo extraer contenido")

        # Problemas de formato
        if analysis_data.get('content_type') == 'unknown':
            issues.append("❓ Tipo de contenido no reconocido - puede afectar la conversión")

        return issues

    def _create_fallback_analysis(self, file_path: Path, error: str) -> ContentAnalysis:
        """Crear análisis básico cuando falla el análisis completo"""
        file_size = file_path.stat().st_size / (1024 * 1024)

        return ContentAnalysis(
            file_type=file_path.suffix.lower().lstrip('.'),
            content_type='unknown',
            complexity_score=50.0,
            quality_score=30.0,
            size_mb=file_size,
            page_count=1,
            word_count=0,
            image_count=0,
            table_count=0,
            has_forms=False,
            has_hyperlinks=False,
            has_embedded_media=False,
            text_to_image_ratio=1.0,
            recommended_formats=[
                {'format': 'pdf', 'reason': 'Formato universal y confiable', 'priority': 'medium'}
            ],
            optimization_suggestions=[
                "⚠️ No se pudo analizar completamente el archivo",
                "🔍 Verifica que el archivo no esté corrupto"
            ],
            quality_issues=[f"❌ Error de análisis: {error}"],
            metadata={'error': error}
        )
