#!/usr/bin/env python3
# ================================
# ANCLORA NEXUS - COMPREHENSIVE CONVERSION TESTS
# Pruebas exhaustivas de todas las conversiones
# ================================

import os
import sys
import time
import tempfile
import requests
import json
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from fpdf import FPDF
import base64

# Configuración
API_BASE_URL = "http://localhost:8000/api/conversion"
TEST_FILES_DIR = "test_files"

def create_test_files_directory():
    """Crea el directorio de archivos de prueba"""
    if not os.path.exists(TEST_FILES_DIR):
        os.makedirs(TEST_FILES_DIR)
    return TEST_FILES_DIR

def create_test_txt_file():
    """Crea archivo TXT de prueba"""
    content = """# Documento de Prueba Anclora Nexus

Este es un archivo de texto de prueba para verificar las conversiones.

## Características del texto:
- Contiene caracteres especiales: áéíóú ñ ¿¡
- Emojis: 🚀 🎉 ✅ ❌ 🔥
- Números: 123456789
- Símbolos: @#$%^&*()_+-={}[]|\\:";'<>?,./

### Lista de elementos:
1. Primer elemento
2. Segundo elemento
3. Tercer elemento

> Cita importante: "La calidad es más importante que la velocidad"

**Texto en negrita** y *texto en cursiva*

Párrafo final con texto normal para completar el documento de prueba.
"""
    
    file_path = os.path.join(TEST_FILES_DIR, "test_document.txt")
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return file_path

def create_test_md_file():
    """Crea archivo Markdown de prueba"""
    content = """# Documento Markdown de Prueba

Este es un **documento Markdown** con diferentes elementos para probar conversiones.

## Características

- Lista con elementos
- **Texto en negrita**
- *Texto en cursiva*
- `Código inline`

### Código de ejemplo

```python
def hello_world():
    print("¡Hola Anclora Nexus!")
    return True
```

### Tabla de datos

| Formato | Calidad | Velocidad |
|---------|---------|-----------|
| PDF     | Alta    | Media     |
| HTML    | Alta    | Alta      |
| DOCX    | Media   | Media     |

> **Nota importante**: Este es un blockquote de prueba

### Enlaces y referencias

- [Anclora Nexus](https://anclora.com)
- [Documentación](https://docs.anclora.com)

---

**Conclusión**: Markdown es excelente para documentación técnica.
"""
    
    file_path = os.path.join(TEST_FILES_DIR, "test_markdown.md")
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return file_path

def create_test_html_file():
    """Crea archivo HTML de prueba"""
    content = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Documento HTML de Prueba - Anclora Nexus</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        h1 { color: #6366f1; }
        .highlight { background-color: #fef3c7; padding: 10px; }
        .code { background-color: #f3f4f6; padding: 5px; font-family: monospace; }
    </style>
</head>
<body>
    <h1>🚀 Anclora Nexus - Documento HTML de Prueba</h1>
    
    <h2>Características del Motor de Conversión</h2>
    <ul>
        <li><strong>IA Integrada</strong>: Análisis automático de archivos</li>
        <li><strong>Múltiples Formatos</strong>: Soporte para 15+ formatos</li>
        <li><strong>Alta Calidad</strong>: Conversiones optimizadas</li>
        <li><strong>Velocidad</strong>: Procesamiento en menos de 3 segundos</li>
    </ul>
    
    <div class="highlight">
        <h3>💡 Recomendación IA</h3>
        <p>Este archivo HTML contiene estilos CSS y estructura semántica que se preservará durante la conversión.</p>
    </div>
    
    <h3>Código de ejemplo</h3>
    <div class="code">
        function convertFile() {<br>
        &nbsp;&nbsp;console.log("Convirtiendo archivo...");<br>
        &nbsp;&nbsp;return true;<br>
        }
    </div>
    
    <h3>Tabla de formatos soportados</h3>
    <table border="1" style="border-collapse: collapse; width: 100%;">
        <tr>
            <th>Formato Origen</th>
            <th>Formatos Destino</th>
            <th>Calidad IA</th>
        </tr>
        <tr>
            <td>TXT</td>
            <td>PDF, HTML, DOCX, MD</td>
            <td>95%</td>
        </tr>
        <tr>
            <td>PDF</td>
            <td>TXT, JPG, PNG</td>
            <td>88%</td>
        </tr>
        <tr>
            <td>MD</td>
            <td>HTML, PDF, DOCX</td>
            <td>92%</td>
        </tr>
    </table>
    
    <footer>
        <p><em>© 2025 Anclora Nexus - Motor de Conversión Inteligente</em></p>
    </footer>
</body>
</html>"""
    
    file_path = os.path.join(TEST_FILES_DIR, "test_webpage.html")
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return file_path

def create_test_docx_file():
    """Crea archivo DOCX de prueba"""
    doc = Document()
    
    # Título
    title = doc.add_heading('Documento DOCX de Prueba - Anclora Nexus', 0)
    
    # Párrafo introductorio
    intro = doc.add_paragraph('Este es un documento de Microsoft Word creado para probar las capacidades de conversión del motor inteligente de Anclora Nexus.')
    intro.bold = True
    
    # Sección con formato
    doc.add_heading('Características del Motor IA', level=1)
    
    features = [
        'Análisis automático de contenido con IA',
        'Predicción de calidad de conversión',
        'Rutas de conversión optimizadas',
        'Soporte para múltiples formatos',
        'Procesamiento por lotes con ZIP'
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    # Tabla
    doc.add_heading('Estadísticas de Conversión', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Formato'
    hdr_cells[1].text = 'Precisión IA'
    hdr_cells[2].text = 'Velocidad'
    
    data = [
        ('TXT → PDF', '95%', '< 1s'),
        ('MD → HTML', '98%', '< 1s'),
        ('DOCX → PDF', '92%', '< 2s'),
        ('PNG → JPG', '99%', '< 1s')
    ]
    
    for format_name, precision, speed in data:
        row_cells = table.add_row().cells
        row_cells[0].text = format_name
        row_cells[1].text = precision
        row_cells[2].text = speed
    
    # Párrafo final
    doc.add_paragraph('\nEste documento contiene formato rico que será preservado durante la conversión inteligente.')
    
    file_path = os.path.join(TEST_FILES_DIR, "test_document.docx")
    doc.save(file_path)
    return file_path

def create_test_image_png():
    """Crea imagen PNG de prueba"""
    # Crear imagen de 800x600 con contenido
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)
    
    # Fondo degradado
    for y in range(600):
        color_value = int(255 * (1 - y / 600))
        draw.line([(0, y), (800, y)], fill=(color_value, color_value + 50, 255))
    
    # Texto principal
    try:
        # Intentar usar una fuente del sistema
        font_large = ImageFont.truetype("arial.ttf", 48)
        font_medium = ImageFont.truetype("arial.ttf", 24)
        font_small = ImageFont.truetype("arial.ttf", 16)
    except:
        # Fallback a fuente por defecto
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    # Título
    draw.text((50, 50), "🚀 ANCLORA NEXUS", fill='white', font=font_large)
    draw.text((50, 120), "Motor de Conversión IA", fill='white', font=font_medium)
    
    # Información
    info_text = [
        "✅ Análisis automático con IA",
        "⚡ Conversión ultra-rápida",
        "🎯 Predicción de calidad",
        "📦 Descarga por lotes",
        "🔒 100% seguro y privado"
    ]
    
    for i, text in enumerate(info_text):
        draw.text((50, 200 + i * 40), text, fill='white', font=font_small)
    
    # Formas geométricas
    draw.rectangle([600, 100, 750, 200], outline='white', width=3)
    draw.ellipse([600, 250, 750, 350], outline='white', width=3)
    
    file_path = os.path.join(TEST_FILES_DIR, "test_image.png")
    img.save(file_path, 'PNG')
    return file_path

def create_test_image_jpg():
    """Crea imagen JPG de prueba"""
    img = Image.new('RGB', (640, 480), color='lightblue')
    draw = ImageDraw.Draw(img)
    
    # Patrón de prueba
    for i in range(0, 640, 40):
        draw.line([(i, 0), (i, 480)], fill='darkblue', width=2)
    
    for i in range(0, 480, 40):
        draw.line([(0, i), (640, i)], fill='darkblue', width=2)
    
    # Texto central
    draw.text((200, 220), "JPG Test Image", fill='darkblue')
    draw.text((180, 250), "Anclora Nexus Engine", fill='darkblue')
    
    file_path = os.path.join(TEST_FILES_DIR, "test_photo.jpg")
    img.save(file_path, 'JPEG', quality=85)
    return file_path

def create_corrupted_files():
    """Crea archivos corruptos para probar manejo de errores"""
    corrupted_files = []
    
    # Archivo TXT con encoding incorrecto
    corrupted_txt = os.path.join(TEST_FILES_DIR, "corrupted_encoding.txt")
    with open(corrupted_txt, 'wb') as f:
        f.write(b'\xff\xfe\x00\x00Invalid UTF-8 content \x80\x81\x82')
    corrupted_files.append(corrupted_txt)
    
    # Archivo PNG corrupto (header incorrecto)
    corrupted_png = os.path.join(TEST_FILES_DIR, "corrupted_image.png")
    with open(corrupted_png, 'wb') as f:
        f.write(b'INVALID_PNG_HEADER' + b'\x00' * 100)
    corrupted_files.append(corrupted_png)
    
    # Archivo PDF corrupto
    corrupted_pdf = os.path.join(TEST_FILES_DIR, "corrupted_document.pdf")
    with open(corrupted_pdf, 'wb') as f:
        f.write(b'%PDF-1.4\nINVALID_PDF_CONTENT\n%%EOF')
    corrupted_files.append(corrupted_pdf)
    
    # Archivo con extensión incorrecta
    fake_docx = os.path.join(TEST_FILES_DIR, "fake_document.docx")
    with open(fake_docx, 'w', encoding='utf-8') as f:
        f.write("Este es un archivo TXT con extensión DOCX")
    corrupted_files.append(fake_docx)
    
    return corrupted_files

def create_edge_case_files():
    """Crea archivos con casos extremos"""
    edge_files = []
    
    # Archivo muy pequeño
    tiny_txt = os.path.join(TEST_FILES_DIR, "tiny_file.txt")
    with open(tiny_txt, 'w', encoding='utf-8') as f:
        f.write("x")
    edge_files.append(tiny_txt)
    
    # Archivo con caracteres especiales en nombre
    special_name = os.path.join(TEST_FILES_DIR, "archivo_con_ñ_y_acentos_áéíóú.txt")
    with open(special_name, 'w', encoding='utf-8') as f:
        f.write("Archivo con nombre especial")
    edge_files.append(special_name)
    
    # Archivo con solo espacios
    spaces_txt = os.path.join(TEST_FILES_DIR, "only_spaces.txt")
    with open(spaces_txt, 'w', encoding='utf-8') as f:
        f.write("   \n\n   \t\t   \n   ")
    edge_files.append(spaces_txt)
    
    # Archivo con líneas muy largas
    long_lines_txt = os.path.join(TEST_FILES_DIR, "long_lines.txt")
    with open(long_lines_txt, 'w', encoding='utf-8') as f:
        long_line = "Esta es una línea extremadamente larga que contiene muchos caracteres para probar cómo el motor de conversión maneja líneas largas sin saltos de línea. " * 50
        f.write(long_line)
    edge_files.append(long_lines_txt)
    
    return edge_files

def test_file_conversion(file_path, target_format, test_name):
    """Prueba la conversión de un archivo específico"""
    print(f"\n🔄 Probando: {test_name}")
    print(f"   📁 Archivo: {os.path.basename(file_path)}")
    print(f"   🎯 Formato destino: {target_format.upper()}")
    
    try:
        # Verificar que el archivo existe
        if not os.path.exists(file_path):
            print(f"   ❌ Archivo no encontrado: {file_path}")
            return False
        
        file_size = os.path.getsize(file_path)
        print(f"   📏 Tamaño: {file_size / 1024:.1f} KB")
        
        # Realizar conversión
        with open(file_path, 'rb') as f:
            files = {'file': (os.path.basename(file_path), f, 'application/octet-stream')}
            data = {'target_format': target_format}
            
            start_time = time.time()
            response = requests.post(f"{API_BASE_URL}/guest-convert", files=files, data=data)
            conversion_time = time.time() - start_time
            
            print(f"   ⏱️  Tiempo: {conversion_time:.2f}s")
            print(f"   📊 Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                if result.get('success'):
                    print(f"   ✅ Conversión exitosa")
                    print(f"   📄 Archivo generado: {result.get('output_filename', 'N/A')}")
                    print(f"   💬 Mensaje: {result.get('message', 'N/A')}")
                    
                    # Verificar descarga
                    download_url = result.get('download_url')
                    if download_url:
                        download_response = requests.get(f"http://localhost:8000{download_url}")
                        if download_response.status_code == 200:
                            output_size = len(download_response.content)
                            print(f"   📦 Descarga exitosa: {output_size / 1024:.1f} KB")
                            return True
                        else:
                            print(f"   ❌ Error en descarga: {download_response.status_code}")
                            return False
                    else:
                        print(f"   ⚠️  Sin URL de descarga")
                        return True
                else:
                    print(f"   ❌ Conversión fallida: {result.get('error', 'Error desconocido')}")
                    return False
            else:
                try:
                    error_data = response.json()
                    print(f"   ❌ Error HTTP: {error_data.get('error', 'Error desconocido')}")
                except:
                    print(f"   ❌ Error HTTP: {response.status_code} - {response.text[:100]}")
                return False
                
    except Exception as e:
        print(f"   ❌ Excepción: {str(e)}")
        return False

def test_ai_analysis(file_path, test_name):
    """Prueba el análisis IA de un archivo"""
    print(f"\n🧠 Análisis IA: {test_name}")
    
    try:
        with open(file_path, 'rb') as f:
            files = {'file': (os.path.basename(file_path), f, 'application/octet-stream')}
            
            start_time = time.time()
            response = requests.post(f"{API_BASE_URL}/ai-analyze", files=files)
            analysis_time = time.time() - start_time
            
            print(f"   ⏱️  Tiempo análisis: {analysis_time:.2f}s")
            
            if response.status_code == 200:
                result = response.json()
                if result.get('success'):
                    analysis = result.get('analysis', {})
                    recommendations = result.get('smart_recommendations', [])
                    
                    print(f"   ✅ Análisis exitoso")
                    print(f"   📊 Tipo: {analysis.get('content_type', 'N/A')}")
                    print(f"   🔍 Complejidad: {analysis.get('complexity_level', 'N/A')}")
                    print(f"   🛡️  Seguridad: {analysis.get('security_level', 'N/A')}")
                    print(f"   💡 Recomendaciones: {len(recommendations)}")
                    
                    if recommendations:
                        best_rec = recommendations[0]
                        print(f"   🎯 Mejor formato: {best_rec.get('target_format', 'N/A').upper()}")
                        print(f"   📈 Calidad predicha: {best_rec.get('quality_prediction', 0)}%")
                    
                    return True
                else:
                    print(f"   ❌ Análisis fallido: {result.get('error', 'Error desconocido')}")
                    return False
            else:
                print(f"   ❌ Error HTTP: {response.status_code}")
                return False
                
    except Exception as e:
        print(f"   ❌ Excepción en análisis: {str(e)}")
        return False

def main():
    """Función principal de pruebas exhaustivas"""
    
    print("🚀 ANCLORA NEXUS - PRUEBAS EXHAUSTIVAS DE CONVERSIÓN")
    print("=" * 70)
    print("Verificando TODAS las conversiones con archivos reales...")
    print()
    
    # Crear directorio de pruebas
    test_dir = create_test_files_directory()
    print(f"📁 Directorio de pruebas: {test_dir}")
    
    # Crear archivos de prueba
    print("\n📝 CREANDO ARCHIVOS DE PRUEBA...")
    test_files = []
    
    try:
        test_files.append(("TXT", create_test_txt_file()))
        test_files.append(("MD", create_test_md_file()))
        test_files.append(("HTML", create_test_html_file()))
        test_files.append(("DOCX", create_test_docx_file()))
        test_files.append(("PNG", create_test_image_png()))
        test_files.append(("JPG", create_test_image_jpg()))
        
        print(f"✅ {len(test_files)} archivos de prueba creados")
        
        # Crear archivos corruptos
        print("\n💀 CREANDO ARCHIVOS CORRUPTOS...")
        corrupted_files = create_corrupted_files()
        print(f"⚠️  {len(corrupted_files)} archivos corruptos creados")
        
        # Crear casos extremos
        print("\n🎭 CREANDO CASOS EXTREMOS...")
        edge_files = create_edge_case_files()
        print(f"🎪 {len(edge_files)} casos extremos creados")
        
    except Exception as e:
        print(f"❌ Error creando archivos: {e}")
        return
    
    # Definir matriz de conversiones a probar
    conversion_matrix = [
        # Conversiones de texto
        ("TXT", ["pdf", "html", "docx", "md"]),
        ("MD", ["html", "pdf", "txt", "docx"]),
        ("HTML", ["pdf", "txt", "md"]),
        ("DOCX", ["pdf", "txt", "html"]),
        
        # Conversiones de imagen
        ("PNG", ["jpg", "gif", "pdf"]),
        ("JPG", ["png", "gif", "pdf"]),
    ]
    
    # Estadísticas
    total_tests = 0
    successful_tests = 0
    failed_tests = 0
    ai_analysis_success = 0
    
    print("\n" + "=" * 70)
    print("🧪 INICIANDO PRUEBAS DE CONVERSIÓN EXHAUSTIVAS")
    print("=" * 70)
    
    # Probar conversiones normales
    for source_format, target_formats in conversion_matrix:
        source_file = next((f[1] for f in test_files if f[0] == source_format), None)
        if not source_file:
            continue
            
        print(f"\n📋 PROBANDO CONVERSIONES DESDE {source_format}")
        print("-" * 50)
        
        # Análisis IA primero
        ai_success = test_ai_analysis(source_file, f"{source_format} Analysis")
        if ai_success:
            ai_analysis_success += 1
        
        # Probar cada conversión
        for target_format in target_formats:
            total_tests += 1
            test_name = f"{source_format} → {target_format.upper()}"
            
            success = test_file_conversion(source_file, target_format, test_name)
            if success:
                successful_tests += 1
            else:
                failed_tests += 1
    
    # Probar archivos corruptos
    print(f"\n💀 PROBANDO ARCHIVOS CORRUPTOS")
    print("-" * 50)
    
    for corrupted_file in corrupted_files:
        total_tests += 1
        file_name = os.path.basename(corrupted_file)
        
        success = test_file_conversion(corrupted_file, "pdf", f"Corrupted: {file_name}")
        if success:
            successful_tests += 1
            print(f"   🔧 ¡Motor IA reparó archivo corrupto!")
        else:
            failed_tests += 1
            print(f"   ✅ Manejo correcto de archivo corrupto")
    
    # Probar casos extremos
    print(f"\n🎭 PROBANDO CASOS EXTREMOS")
    print("-" * 50)
    
    for edge_file in edge_files:
        total_tests += 1
        file_name = os.path.basename(edge_file)
        
        success = test_file_conversion(edge_file, "pdf", f"Edge case: {file_name}")
        if success:
            successful_tests += 1
        else:
            failed_tests += 1
    
    # Resumen final
    print("\n" + "=" * 70)
    print("📊 RESUMEN DE PRUEBAS EXHAUSTIVAS")
    print("=" * 70)
    
    success_rate = (successful_tests / total_tests * 100) if total_tests > 0 else 0
    ai_success_rate = (ai_analysis_success / len(test_files) * 100) if test_files else 0
    
    print(f"📈 Pruebas totales: {total_tests}")
    print(f"✅ Conversiones exitosas: {successful_tests}")
    print(f"❌ Conversiones fallidas: {failed_tests}")
    print(f"📊 Tasa de éxito: {success_rate:.1f}%")
    print(f"🧠 Análisis IA exitosos: {ai_analysis_success}/{len(test_files)} ({ai_success_rate:.1f}%)")
    
    print(f"\n🎯 EVALUACIÓN DEL SISTEMA:")
    if success_rate >= 90:
        print("🏆 EXCELENTE: Sistema funcionando perfectamente")
    elif success_rate >= 75:
        print("✅ BUENO: Sistema funcionando bien con algunos problemas menores")
    elif success_rate >= 50:
        print("⚠️  REGULAR: Sistema necesita mejoras")
    else:
        print("❌ CRÍTICO: Sistema requiere revisión urgente")
    
    print(f"\n💡 RECOMENDACIONES:")
    if failed_tests > 0:
        print("   • Revisar conversiones fallidas")
        print("   • Mejorar manejo de errores")
        print("   • Optimizar algoritmos de conversión")
    
    if ai_analysis_success < len(test_files):
        print("   • Verificar servicios de IA")
        print("   • Configurar API keys si es necesario")
    
    print("   • Probar con archivos más grandes")
    print("   • Implementar más formatos si es necesario")
    
    # Limpiar archivos de prueba
    print(f"\n🧹 Limpiando archivos de prueba...")
    try:
        import shutil
        shutil.rmtree(TEST_FILES_DIR)
        print("✅ Archivos de prueba eliminados")
    except Exception as e:
        print(f"⚠️  Error limpiando: {e}")

if __name__ == "__main__":
    main()
